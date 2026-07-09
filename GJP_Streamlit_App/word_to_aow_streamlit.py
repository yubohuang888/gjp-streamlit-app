import io
import re
from dataclasses import dataclass
from typing import Dict, Iterable, List, Optional, Sequence, Set, Tuple

import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

try:
    import streamlit as st
except ModuleNotFoundError:
    st = None


RESPONSIBILITY_HEADERS = ("area of work", "no", "responsibilities")
KSA_HEADER_KEYWORDS = ("knowledge", "skills", "abilities")


@dataclass
class Responsibility:
    theme: str
    number: str
    text: str


@dataclass
class ConversionResult:
    dataframe: pd.DataFrame
    workbook_bytes: bytes
    responsibilities: List[Responsibility]
    ksas: List[str]
    warnings: List[str]


def clean_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\u00a0", " ")
    return re.sub(r"\s+", " ", text).strip()


def cell_text(cell) -> str:
    return clean_text(cell.text)


def row_texts(row) -> List[str]:
    return [cell_text(cell) for cell in row.cells]


def extract_numbers(text: str) -> Set[str]:
    return set(re.findall(r"\d+", clean_text(text)))


def looks_like_responsibility_table(table) -> bool:
    if not table.rows:
        return False
    header = " | ".join(row_texts(table.rows[0])).lower()
    return all(keyword in header for keyword in RESPONSIBILITY_HEADERS)


def looks_like_ksa_table(table) -> bool:
    if not table.rows:
        return False
    first_cell = cell_text(table.rows[0].cells[0]).lower()
    return all(keyword in first_cell for keyword in KSA_HEADER_KEYWORDS)


def find_responsibility_table(document: Document):
    candidates = [table for table in document.tables if looks_like_responsibility_table(table)]
    if candidates:
        return candidates[0]
    raise ValueError("No responsibilities table was found. The Word file should include Area of Work, No., and Responsibilities headers.")


def find_ksa_table(document: Document):
    candidates = [table for table in document.tables if looks_like_ksa_table(table)]
    if candidates:
        return candidates[0]
    raise ValueError("No KSA mapping table was found. The Word file should include a Knowledge, Skills, and Abilities table.")


def header_index(headers: Sequence[str], keywords: Iterable[str]) -> Optional[int]:
    lowered = [header.lower() for header in headers]
    for keyword in keywords:
        for index, header in enumerate(lowered):
            if keyword in header:
                return index
    return None


def parse_responsibilities(document: Document) -> List[Responsibility]:
    table = find_responsibility_table(document)
    headers = row_texts(table.rows[0])
    theme_col = header_index(headers, ("area of work", "theme"))
    number_col = header_index(headers, ("no.", "no", "nr", "number"))
    resp_col = header_index(headers, ("responsibilities", "responsibility"))

    if theme_col is None or number_col is None or resp_col is None:
        raise ValueError("The responsibilities table headers could not be recognized. Required headers: Area of Work, No., Responsibilities.")

    responsibilities: List[Responsibility] = []
    last_theme = ""
    for row in table.rows[1:]:
        cells = row_texts(row)
        if len(cells) <= max(theme_col, number_col, resp_col):
            continue
        theme = cells[theme_col] or last_theme
        number_match = re.search(r"\d+", cells[number_col])
        number = number_match.group(0) if number_match else clean_text(cells[number_col])
        text = cells[resp_col]
        if theme:
            last_theme = theme
        if number and text:
            responsibilities.append(Responsibility(theme=theme, number=number, text=text))

    if not responsibilities:
        raise ValueError("No valid responsibility rows were found.")
    return responsibilities


def parse_ksa_mapping(document: Document) -> Tuple[List[str], Dict[str, Set[str]]]:
    table = find_ksa_table(document)
    ksas: List[str] = []
    mapping_by_number: Dict[str, Set[str]] = {}

    for row in table.rows[1:]:
        cells = row_texts(row)
        if not cells:
            continue
        ksa = cells[0]
        if not ksa or all(keyword in ksa.lower() for keyword in KSA_HEADER_KEYWORDS):
            continue
        if ksa not in ksas:
            ksas.append(ksa)
        for number in extract_numbers(" ".join(cells[1:])):
            mapping_by_number.setdefault(number, set()).add(ksa)

    if not ksas:
        raise ValueError("No valid KSA rows were found.")
    return ksas, mapping_by_number


def build_dataframe(
    responsibilities: List[Responsibility],
    ksas: List[str],
    mapping_by_number: Dict[str, Set[str]],
    mark: str,
    include_assigned_sme: bool,
    collapse_theme: bool,
) -> pd.DataFrame:
    rows = []
    previous_theme = None
    for responsibility in responsibilities:
        shown_theme = responsibility.theme
        if collapse_theme and responsibility.theme == previous_theme:
            shown_theme = ""
        previous_theme = responsibility.theme

        row = {
            "Theme": shown_theme,
            "Number": responsibility.number,
            "Responsibilities": responsibility.text,
        }
        if include_assigned_sme:
            row["Assigned SME"] = ""
        linked_ksas = mapping_by_number.get(responsibility.number, set())
        for ksa in ksas:
            row[ksa] = mark if ksa in linked_ksas else ""
        rows.append(row)

    base_columns = ["Theme", "Number", "Responsibilities"]
    if include_assigned_sme:
        base_columns.append("Assigned SME")
    return pd.DataFrame(rows, columns=base_columns + ksas)


def apply_sheet_style(ws, row_count: int, col_count: int, include_assigned_sme: bool) -> None:
    dark_fill = PatternFill("solid", fgColor="1F4E78")
    note_fill = PatternFill("solid", fgColor="EAF3F8")
    header_font = Font(color="FFFFFF", bold=True)
    note_font = Font(color="1F4E78", bold=True)
    thin_gray = Side(style="thin", color="D9E2EA")
    medium_blue = Side(style="medium", color="1F4E78")
    border = Border(left=thin_gray, right=thin_gray, top=thin_gray, bottom=thin_gray)

    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "E4" if include_assigned_sme else "D4"

    ws.cell(1, 1).value = "Please enter an X/\u00d7 to map KSAs to responsibilities"
    ws.cell(2, 1).value = "Generated from Word responsibilities and KSA mapping tables"
    for row in (1, 2):
        for col in range(1, col_count + 1):
            cell = ws.cell(row, col)
            cell.fill = note_fill
            cell.font = note_font
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 22
    ws.row_dimensions[2].height = 22

    for cell in ws[3]:
        cell.fill = dark_fill
        cell.font = header_font
        cell.border = Border(top=medium_blue, bottom=medium_blue)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[3].height = 58

    theme_fills = [
        PatternFill("solid", fgColor="FFFFFF"),
        PatternFill("solid", fgColor="F6FAFD"),
    ]
    current_fill_index = 0
    last_visible_theme = None
    for row in range(4, row_count + 4):
        visible_theme = clean_text(ws.cell(row, 1).value)
        if visible_theme and visible_theme != last_visible_theme:
            current_fill_index = 1 - current_fill_index
            last_visible_theme = visible_theme
        fill = theme_fills[current_fill_index]
        for col in range(1, col_count + 1):
            cell = ws.cell(row, col)
            cell.fill = fill
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        ws.cell(row, 2).alignment = Alignment(horizontal="center", vertical="top", wrap_text=True)
        first_ksa_col = 5 if include_assigned_sme else 4
        for col in range(first_ksa_col, col_count + 1):
            ws.cell(row, col).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.row_dimensions[row].height = 54

    widths = {
        1: 22,
        2: 8,
        3: 72,
    }
    if include_assigned_sme:
        widths[4] = 16
        first_ksa_col = 5
    else:
        first_ksa_col = 4
    for col in range(first_ksa_col, col_count + 1):
        widths[col] = 24
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = width

    ws.auto_filter.ref = f"A3:{get_column_letter(col_count)}{row_count + 3}"


def dataframe_to_xlsx(df: pd.DataFrame, include_assigned_sme: bool) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "AoW Mapping"

    for col_index, column_name in enumerate(df.columns, start=1):
        ws.cell(3, col_index).value = column_name
    for row_index, values in enumerate(df.itertuples(index=False, name=None), start=4):
        for col_index, value in enumerate(values, start=1):
            ws.cell(row_index, col_index).value = value

    apply_sheet_style(ws, len(df), len(df.columns), include_assigned_sme)
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


def convert_docx(
    file_bytes: bytes,
    mark: str,
    include_assigned_sme: bool,
    collapse_theme: bool,
) -> ConversionResult:
    document = Document(io.BytesIO(file_bytes))
    responsibilities = parse_responsibilities(document)
    ksas, mapping_by_number = parse_ksa_mapping(document)
    df = build_dataframe(
        responsibilities=responsibilities,
        ksas=ksas,
        mapping_by_number=mapping_by_number,
        mark=mark,
        include_assigned_sme=include_assigned_sme,
        collapse_theme=collapse_theme,
    )
    workbook_bytes = dataframe_to_xlsx(df, include_assigned_sme=include_assigned_sme)

    responsibility_numbers = {item.number for item in responsibilities}
    mapped_numbers = set(mapping_by_number)
    warnings = []
    unknown_numbers = sorted(mapped_numbers - responsibility_numbers, key=lambda x: int(x) if x.isdigit() else 99999)
    if unknown_numbers:
        warnings.append(f"The KSA table references responsibility numbers that were not found: {', '.join(unknown_numbers)}")
    unmapped_numbers = sorted(responsibility_numbers - mapped_numbers, key=lambda x: int(x) if x.isdigit() else 99999)
    if unmapped_numbers:
        warnings.append(f"These responsibility numbers do not have any KSA marks: {', '.join(unmapped_numbers)}")

    return ConversionResult(
        dataframe=df,
        workbook_bytes=workbook_bytes,
        responsibilities=responsibilities,
        ksas=ksas,
        warnings=warnings,
    )


def main() -> None:
    if st is None:
        raise RuntimeError(
            "Streamlit is not installed. Run: pip install streamlit python-docx openpyxl pandas"
        )

    st.set_page_config(page_title="Word to AoW/KSA Excel", layout="wide")
    st.markdown(
        """
        <style>
            .stApp {
                background: #f6f8fb;
            }
            .block-container {
                padding-top: 2.25rem;
                padding-bottom: 3rem;
                max-width: 1180px;
            }
            .hero {
                background: linear-gradient(135deg, #123c69 0%, #1f6f8b 56%, #2a9d8f 100%);
                border-radius: 16px;
                padding: 34px 38px;
                color: white;
                box-shadow: 0 18px 45px rgba(18, 60, 105, 0.18);
                margin-bottom: 24px;
            }
            .hero h1 {
                margin: 0;
                font-size: 2.35rem;
                letter-spacing: 0;
            }
            .hero p {
                margin: 12px 0 0;
                font-size: 1.03rem;
                line-height: 1.55;
                opacity: 0.94;
                max-width: 820px;
            }
            .panel {
                background: #ffffff;
                border: 1px solid #e3e8ef;
                border-radius: 12px;
                padding: 22px 24px;
                box-shadow: 0 8px 24px rgba(15, 23, 42, 0.05);
                margin-bottom: 18px;
            }
            .panel h3 {
                margin: 0 0 8px;
                color: #102a43;
                font-size: 1.18rem;
            }
            .panel p {
                color: #526173;
                margin: 0;
                line-height: 1.5;
            }
            div[data-testid="stMetric"] {
                background: #ffffff;
                border: 1px solid #e3e8ef;
                border-radius: 12px;
                padding: 12px 16px;
                box-shadow: 0 6px 18px rgba(15, 23, 42, 0.04);
            }
            section[data-testid="stSidebar"] {
                background: #ffffff;
                border-right: 1px solid #e3e8ef;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="hero">
            <h1>Word to AoW/KSA Excel Converter</h1>
            <p>
                Upload a Word profile with responsibilities and KSA mapping tables.
                The app converts it into an Excel matrix with Theme, Number,
                Responsibilities, and one column for each KSA.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="panel">
            <h3>Upload source document</h3>
            <p>
                The Word file should contain one responsibilities table and one
                Knowledge, Skills, and Abilities mapping table.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.header("Export settings")
        mark = "X"
        st.caption("KSA mark: X")
        include_assigned_sme = st.checkbox("Include blank Assigned SME column", value=True)
        collapse_theme = st.checkbox("Show each theme only once", value=True)

    uploaded_file = st.file_uploader("Choose a .docx file", type=["docx"])

    if uploaded_file is None:
        st.info("Upload a Word document to begin.")
        return

    try:
        result = convert_docx(
            file_bytes=uploaded_file.getvalue(),
            mark=mark,
            include_assigned_sme=include_assigned_sme,
            collapse_theme=collapse_theme,
        )
    except Exception as exc:
        st.error(f"Conversion failed: {exc}")
        st.stop()

    col1, col2, col3 = st.columns(3)
    col1.metric("Responsibilities", len(result.responsibilities))
    col2.metric("KSA columns", len(result.ksas))
    col3.metric("Output columns", len(result.dataframe.columns))

    for warning in result.warnings:
        st.warning(warning)

    st.subheader("Preview")
    st.dataframe(result.dataframe, use_container_width=True, height=520)

    output_name = uploaded_file.name.rsplit(".", 1)[0] + "_AoW_KSA.xlsx"
    st.download_button(
        label="Download Excel",
        data=result.workbook_bytes,
        file_name=output_name,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        type="primary",
    )


if __name__ == "__main__":
    main()
