"""
GJP document generator page for Streamlit.

What it does
------------
1) Accepts either one KSA Excel file or a ZIP containing specialty Excel files.
2) Uploads one GJP Word template.
3) Uploads a Word file containing Work interaction and Introduction text.
4) Generates all configured levels for the department/workstream.
5) Replaces level text, updates work experience years, replaces Work Interactions and Introduction,
   inserts available tables, and removes Responsibility/KSA sections where level data is unavailable.

Install requirements
--------------------
pip install streamlit python-docx openpyxl

How to use in your main GJP_website.py
--------------------------------------
from gjp_document_generator import render_gjp_document_generator
render_gjp_document_generator()
"""

from __future__ import annotations

import io
import re
import zipfile
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

try:
    import streamlit as st
except ModuleNotFoundError:
    st = None
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet


# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------

OUTPUT_LEVELS = ["P-1", "P-2", "P-3", "P-4", "P-5", "D-1", "NO-A", "NO-B", "NO-C", "NO-D"]

# NO levels may reuse P-level narrative text, but never Responsibility/KSA tables.
LEVEL_EQUIVALENT = {
    "NO-A": "P-1",
    "NO-B": "P-2",
    "NO-C": "P-3",
    "NO-D": "P-4",
}

# Work experience wording requested by the user.
WORK_EXPERIENCE_YEARS = {
    "P-1": None,      # keep the template wording
    "NO-A": None,     # keep the template wording
    "P-2": "two",
    "NO-B": "two",
    "P-3": "five",
    "NO-C": "five",
    "P-4": "seven",
    "NO-D": "seven",
    "P-5": "ten",
    "D-1": "fifteen",
}

# Theme colors learned from the uploaded 6.2 Transposing KSA mapping template.
# The colors are used in both the Responsibilities table and the KSA mapping table.
THEME_COLORS = {
    "Risk Assessment & Security Planning": "D9D9D9",
    "Stakeholder Engagement & Coordination": "FFF099",
    "Crisis Management & Response": "A9D08E",
    "Security Operations & Compliance": "F8CBAD",
    "Capacity Building & Training": "AEAAAA",
    "Reporting & Information Management": "CC99FF",
    "Peace Operations & Mission Support": "92D050",
    "Managerial": "C65911",
}
DEFAULT_THEME_COLOR = "D9D9D9"
HEADER_FILL = "000000"
HEADER_FONT = "FFFFFF"


@dataclass
class Responsibility:
    number: int
    theme: str
    text: str
    source_code: str = ""
    mandatory: bool = False
    fill_hex: str = ""
    theme_fill_hex: str = ""
    number_fill_hex: str = ""
    text_fill_hex: str = ""
    theme_bold: bool = False
    number_bold: bool = False
    text_bold: bool = False


@dataclass
class KsaRow:
    number: int
    text: str
    responsibilities: List[int] = field(default_factory=list)
    text_bold: bool = False
    text_fill_hex: str = ""
    mapping_fills: Dict[int, str] = field(default_factory=dict)
    mapping_bolds: Dict[int, bool] = field(default_factory=dict)


@dataclass
class LevelTables:
    level: str
    specialty: str
    responsibilities: List[Responsibility]
    ksa_rows: List[KsaRow]
    header_title: str = ""


@dataclass
class JobInfo:
    level: str
    file_name: str
    job_title: str
    ccog_code: str
    job_code: str


# -----------------------------------------------------------------------------
# Basic Word helpers
# -----------------------------------------------------------------------------

def _normalize_level(level: str) -> str:
    s = str(level).strip().upper().replace("–", "-").replace("—", "-").replace("â€“", "-").replace("â€”", "-")
    # Convert compact forms from Excel, such as P2/D1/NOA, to P-2/D-1/NO-A.
    m = re.fullmatch(r"P\s*-?\s*([1-5])", s)
    if m:
        return f"P-{m.group(1)}"
    m = re.fullmatch(r"D\s*-?\s*1", s)
    if m:
        return "D-1"
    m = re.fullmatch(r"NO\s*-?\s*([A-D])", s)
    if m:
        return f"NO-{m.group(1)}"
    return s


def _compact_level(level: str) -> str:
    return _normalize_level(level).replace("-", "")


def _para_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    """Replace text in a paragraph while keeping paragraph style. Run-level style may be flattened."""
    text = _para_text(paragraph)
    new_text = text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text != text:
        # Preserve the style of the first run when possible.
        if paragraph.runs:
            first_run = paragraph.runs[0]
            for r in list(paragraph.runs):
                r.text = ""
            first_run.text = new_text
        else:
            paragraph.add_run(new_text)


def replace_level_everywhere(doc: DocumentObject, source_level: str, target_level: str) -> None:
    """Replace the template level everywhere: body paragraphs, tables, headers and footers."""
    levels_to_replace = {source_level, _normalize_level(source_level), _compact_level(source_level)}
    replacements = {}
    for old in levels_to_replace:
        replacements[old] = target_level if "-" in old else _compact_level(target_level)

    containers = [doc]
    for section in doc.sections:
        containers.extend([section.header, section.footer])

    for container in containers:
        for paragraph in container.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)


def _replace_everywhere(doc: DocumentObject, replacements: Dict[str, str]) -> None:
    containers = [doc]
    for section in doc.sections:
        containers.extend([section.header, section.footer])

    for container in containers:
        for paragraph in container.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)


def detect_template_level(doc: DocumentObject) -> str:
    text = "\n".join(p.text for p in doc.paragraphs[:20])
    m = re.search(r"\b(P-\d|D-1|NO-[A-D])\b", text, flags=re.I)
    return _normalize_level(m.group(1)) if m else "P-2"


def replace_work_experience_years(doc: DocumentObject, target_level: str) -> None:
    if target_level in {"P-1", "NO-A"}:
        _insert_p1_no_experience_sentence(doc)
        return

    year_word = WORK_EXPERIENCE_YEARS.get(target_level)
    if not year_word:
        return

    pattern = re.compile(r"(A minimum of\s+)(one|two|three|four|five|six|seven|eight|nine|ten|fifteen|\d+)(\s+years)", re.I)

    def replace_years_in_paragraph(paragraph) -> None:
        text = _para_text(paragraph)
        new_text = pattern.sub(lambda m: f"{m.group(1)}{year_word}{m.group(3)}", text)
        if new_text != text:
            if paragraph.runs:
                first = paragraph.runs[0]
                for r in paragraph.runs:
                    r.text = ""
                first.text = new_text
            else:
                paragraph.add_run(new_text)

    for paragraph in doc.paragraphs:
        replace_years_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_years_in_paragraph(paragraph)


def _insert_p1_no_experience_sentence(doc: DocumentObject) -> None:
    sentence = "No experience required with advanced university degree."
    body = doc._body._element
    children = list(body)
    paragraph_by_xml = {p._p: p for p in doc.paragraphs}
    target = None
    for child in children:
        paragraph = paragraph_by_xml.get(child)
        if paragraph is None:
            continue
        text = paragraph.text.strip()
        if (
            "The work experience requirement is the same for all job openings" in text
            and "positions based on this Generic Job Profile" in text
        ):
            target = paragraph
            break
    if target is None:
        return

    target_index = children.index(target._p)
    for child in children[target_index + 1:target_index + 4]:
        paragraph = paragraph_by_xml.get(child)
        if paragraph is not None and sentence.lower() in paragraph.text.lower():
            return

    new_p = _insert_paragraph_after(target, sentence)
    for run in new_p.runs:
        run.font.size = Pt(11)


# -----------------------------------------------------------------------------
# Parse job title / job code workbook
# -----------------------------------------------------------------------------

def _level_from_text(value: str) -> str:
    match = re.search(r"\b(P\s*-?\s*[1-5]|D\s*-?\s*1|NO\s*-?\s*[A-D])\b", str(value or ""), flags=re.I)
    return _normalize_level(match.group(1)) if match else ""


def parse_job_info_workbook(path: str | Path) -> Dict[str, JobInfo]:
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    header_row = None
    header_map: Dict[str, int] = {}

    for row in range(1, min(ws.max_row, 20) + 1):
        values = {
            str(ws.cell(row, col).value or "").strip().lower(): col
            for col in range(1, ws.max_column + 1)
            if ws.cell(row, col).value is not None
        }
        if {"file name", "job title", "ccog code", "job code"}.issubset(values):
            header_row = row
            header_map = values
            break

    if header_row is None:
        raise ValueError("The job info workbook must include File Name, Job Title, CCOG Code, and Job Code columns.")

    result: Dict[str, JobInfo] = {}
    for row in range(header_row + 1, ws.max_row + 1):
        file_name = _cell_text(ws, row, header_map["file name"])
        job_title = _cell_text(ws, row, header_map["job title"])
        ccog_code = _cell_text(ws, row, header_map["ccog code"])
        job_code = _cell_text(ws, row, header_map["job code"])
        level = _level_from_text(job_title) or _level_from_text(file_name)
        if not level:
            continue
        result[level] = JobInfo(
            level=level,
            file_name=file_name or f"GJP 2.0 - {level}.docx",
            job_title=job_title,
            ccog_code=ccog_code,
            job_code=job_code,
        )

    if not result:
        raise ValueError("No level rows were found in the job info workbook.")
    return result


def _strip_level_from_title(title: str, level: str) -> str:
    text = str(title or "").strip()
    if not text:
        return ""
    text = re.sub(rf"\s*\b{re.escape(level)}\b\s*$", "", text, flags=re.I)
    return re.sub(r"\s+", " ", text).strip()


def _replace_job_info(
    doc: DocumentObject,
    source_level: str,
    target_level: str,
    job_info: Optional[JobInfo],
    source_job_info: Optional[JobInfo],
) -> None:
    if not job_info:
        return

    replacements: Dict[str, str] = {}
    if source_job_info and source_job_info.job_title and job_info.job_title:
        source_title = source_job_info.job_title
        target_title = job_info.job_title
        replacements[source_title] = target_title
        source_title_no_level = _strip_level_from_title(source_title, source_level)
        target_title_no_level = _strip_level_from_title(target_title, target_level)
        if source_title_no_level and target_title_no_level and source_title_no_level != target_title_no_level:
            replacements[source_title_no_level] = target_title_no_level

    if replacements:
        _replace_everywhere(doc, replacements)

    ccog_pattern = re.compile(r"(The CCOG Code is\s+)(.*?)(\s+and the Job Code is\s+)(.*?)(\.)", re.I)
    for paragraph in doc.paragraphs:
        text = _para_text(paragraph)
        new_text = ccog_pattern.sub(
            lambda m: f"{m.group(1)}{job_info.ccog_code}{m.group(3)}{job_info.job_code}{m.group(5)}",
            text,
        )
        if new_text != text:
            _replace_in_paragraph(paragraph, {text: new_text})


# -----------------------------------------------------------------------------
# Parse Work interaction / Introduction Word file
# -----------------------------------------------------------------------------

def parse_intro_interaction_docx(path: str | Path) -> Tuple[Dict[str, str], Dict[str, Dict[str, str]]]:
    """Return interactions[level] and introductions[level][specialty]."""
    doc = Document(str(path))
    paragraphs = [p.text.strip().replace("[Done]", "").strip() for p in doc.paragraphs if p.text.strip()]

    line_re = re.compile(
        r"^\s*(P\s*-?\s*[1-5]|D\s*-?\s*1|NO\s*-?\s*[A-D])\s*-\s*(.*?)\s*-\s*(Work\s*Interactions?|Introduction)\s*:\s*(.*)$",
        re.I,
    )
    interactions: Dict[str, str] = {}
    introductions: Dict[str, Dict[str, str]] = {}
    parsed_line_format = False
    for paragraph_text in paragraphs:
        match = line_re.match(paragraph_text)
        if not match:
            continue
        parsed_line_format = True
        level = _normalize_level(match.group(1))
        specialty = re.sub(r"\s+", " ", match.group(2)).strip()
        kind = match.group(3).lower()
        body = re.sub(r"\s+", " ", match.group(4)).strip()
        if not body:
            continue
        if kind.startswith("work"):
            interactions[level] = body
        else:
            introductions.setdefault(level, {})[specialty] = body

    if parsed_line_format:
        return interactions, introductions

    text = "\n".join(paragraphs)

    # Split into two sections if headings exist.
    inter_part = text
    intro_part = text
    m_intro = re.search(r"\bIntroduction\s*:", text, flags=re.I)
    if m_intro:
        inter_part = text[:m_intro.start()]
        intro_part = text[m_intro.end():]
    m_inter = re.search(r"\bWork\s*interaction\s*:", inter_part, flags=re.I)
    if m_inter:
        inter_part = inter_part[m_inter.end():]

    def extract_by_level(section_text: str) -> Dict[str, str]:
        result: Dict[str, str] = {}
        # Matches P-2:, P-3:, P-4:, P-5:, D-1:, NO-A: etc.
        label_re = re.compile(r"(?im)(?:^|\n)\s*(P-\d|D-1|NO-[A-D])\s*:\s*")
        matches = list(label_re.finditer(section_text))
        for i, m in enumerate(matches):
            level = _normalize_level(m.group(1))
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(section_text)
            body = section_text[start:end].strip()
            body = re.sub(r"\s+", " ", body)
            if body:
                result[level] = body
        return result

    interactions = extract_by_level(inter_part)
    introductions = {
        level: {"": body}
        for level, body in extract_by_level(intro_part).items()
    }
    return interactions, introductions


# -----------------------------------------------------------------------------
# Parse the level-differentiated Excel workbook
# -----------------------------------------------------------------------------

def _cell_text(ws: Worksheet, row: int, col: int) -> str:
    v = ws.cell(row, col).value
    return "" if v is None else str(v).strip()


def _find_header_rows(ws: Worksheet) -> List[int]:
    rows = []
    for r in range(1, ws.max_row + 1):
        row_values = [_cell_text(ws, r, c).lower().replace("\n", " ") for c in range(1, min(ws.max_column, 12) + 1)]
        if "theme" in row_values and "level" in row_values and any("responsibility" in x for x in row_values):
            rows.append(r)
    return rows


def _header_map(ws: Worksheet, header_row: int) -> Dict[str, int]:
    mapping: Dict[str, int] = {}
    for c in range(1, ws.max_column + 1):
        value = _cell_text(ws, header_row, c).lower().replace("\n", " ")
        if value == "theme":
            mapping["theme"] = c
        elif value == "level":
            mapping["level"] = c
        elif "resp" in value and "code" in value:
            mapping["code"] = c
        elif value == "responsibility" or "responsibility" in value:
            mapping["responsibility"] = c
        elif "standard" in value or "mandatory" in value or "grouping" in value:
            mapping["mandatory"] = c
    return mapping


def _detect_specialty_name(ws: Worksheet, header_row: int) -> str:
    # If the table has a title above it, use it as specialty; otherwise General.
    for r in range(header_row - 1, max(0, header_row - 6), -1):
        values = [_cell_text(ws, r, c) for c in range(1, min(ws.max_column, 8) + 1)]
        joined = " ".join(v for v in values if v)
        if joined and not re.search(r"security coordination officer|level differentiated|mapping", joined, re.I):
            return joined[:80]
    return "General"


def _is_marked(value) -> bool:
    if value is None:
        return False
    s = str(value).strip().lower()
    return s in {"x", "yes", "y", "1", "true", "mandatory", "required"}


def _theme_color(theme: str) -> str:
    if not theme:
        return DEFAULT_THEME_COLOR
    clean = re.sub(r"\s+", " ", theme.strip())
    if clean in THEME_COLORS:
        return THEME_COLORS[clean]
    for known, color in THEME_COLORS.items():
        if clean.lower() == known.lower():
            return color
    return DEFAULT_THEME_COLOR


def parse_level_tables_from_excel(path: str | Path) -> Dict[str, List[LevelTables]]:
    """
    Parse the SCO_KSA_Level_Differentiated workbook.

    The parser supports one or multiple tables per sheet. It looks for rows with headers:
    Theme | Level | Resp. Code | Responsibility | OHR Standard | KSA columns...
    """
    wb = load_workbook(path, data_only=False)
    output: Dict[Tuple[str, str], Dict] = {}

    for ws in wb.worksheets:
        header_rows = _find_header_rows(ws)
        if not header_rows:
            continue
        header_rows.append(ws.max_row + 2)

        for idx, header_row in enumerate(header_rows[:-1]):
            next_header = header_rows[idx + 1]
            hmap = _header_map(ws, header_row)
            if not {"theme", "level", "responsibility"}.issubset(hmap):
                continue

            specialty = _detect_specialty_name(ws, header_row)
            ksa_start = max(hmap.values()) + 1
            # In the uploaded file, KSA headers are in row 4 while table headers are row 2.
            ksa_text_row = header_row + 2 if header_row + 2 <= ws.max_row else header_row
            ksa_cols: List[Tuple[int, str]] = []
            for c in range(ksa_start, ws.max_column + 1):
                ksa_text = _cell_text(ws, ksa_text_row, c)
                if ksa_text and len(ksa_text) > 5:
                    ksa_cols.append((c, ksa_text))

            current_theme = ""
            for r in range(header_row + 3, min(next_header, ws.max_row + 1)):
                level = _normalize_level(_cell_text(ws, r, hmap["level"]))
                responsibility_text = _cell_text(ws, r, hmap["responsibility"])
                if not level or not responsibility_text:
                    continue
                raw_theme = _cell_text(ws, r, hmap["theme"])
                if raw_theme:
                    current_theme = raw_theme
                theme = current_theme or "General"
                code = _cell_text(ws, r, hmap.get("code", 0)) if hmap.get("code") else ""
                mandatory = False
                if hmap.get("mandatory"):
                    mandatory = _is_marked(ws.cell(r, hmap["mandatory"]).value)
                # Bold source responsibility is also treated as mandatory.
                try:
                    mandatory = mandatory or bool(ws.cell(r, hmap["responsibility"]).font.bold)
                except Exception:
                    pass

                key = (level, specialty)
                if key not in output:
                    output[key] = {"responsibilities": [], "ksa_map": {}}
                resp_no = len(output[key]["responsibilities"]) + 1
                output[key]["responsibilities"].append(
                    Responsibility(number=resp_no, theme=theme, text=responsibility_text, source_code=code, mandatory=mandatory)
                )
                for c, ksa_text in ksa_cols:
                    if _is_marked(ws.cell(r, c).value):
                        output[key]["ksa_map"].setdefault(ksa_text, []).append(resp_no)

    result: Dict[str, List[LevelTables]] = {}
    for (level, specialty), data in output.items():
        ksa_rows = [
            KsaRow(number=i + 1, text=ksa_text, responsibilities=resp_numbers)
            for i, (ksa_text, resp_numbers) in enumerate(data["ksa_map"].items())
        ]
        result.setdefault(level, []).append(
            LevelTables(level=level, specialty=specialty, responsibilities=data["responsibilities"], ksa_rows=ksa_rows)
        )
    return result


def get_tables_for_level(all_tables: Dict[str, List[LevelTables]], target_level: str) -> List[LevelTables]:
    """Return only exact uploaded data; never manufacture a missing level."""
    return all_tables.get(target_level, [])


OUTPUT_SHEET_RE = re.compile(
    r"^output[\s_-]*(p[\s_-]?[1-5]|d[\s_-]?1|no[\s_-]?[a-d]).*?[\s_-](resp|ksa)$",
    re.I,
)


def _cell_fill_hex(cell) -> str:
    fill = cell.fill
    if not fill or fill.fill_type != "solid":
        return ""
    color = fill.fgColor
    return _openpyxl_color_to_hex(color, cell.parent.parent)


def _openpyxl_color_to_hex(color, workbook=None) -> str:
    if color is None or color.type == "auto":
        return ""
    if color.type == "rgb" and color.rgb:
        value = str(color.rgb)
        return value[-6:] if len(value) >= 6 else ""
    if color.type == "indexed" and color.indexed is not None:
        try:
            from openpyxl.styles.colors import COLOR_INDEX
            value = COLOR_INDEX[int(color.indexed)]
            return str(value)[-6:]
        except Exception:
            return ""
    if color.type == "theme" and workbook is not None and color.theme is not None:
        value = _theme_color_to_hex(workbook, int(color.theme), float(color.tint or 0))
        return value
    return ""


def _theme_color_to_hex(workbook, theme_index: int, tint: float = 0) -> str:
    try:
        from xml.etree import ElementTree as ET
        theme_xml = workbook.loaded_theme
        if not theme_xml:
            return ""
        root = ET.fromstring(theme_xml)
        ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        scheme = root.find(".//a:clrScheme", ns)
        if scheme is None:
            return ""
        theme_order = [
            "lt1", "dk1", "lt2", "dk2", "accent1", "accent2",
            "accent3", "accent4", "accent5", "accent6", "hlink", "folHlink",
        ]
        if theme_index >= len(theme_order):
            return ""
        node = scheme.find(f"a:{theme_order[theme_index]}", ns)
        if node is None:
            return ""
        srgb = node.find("a:srgbClr", ns)
        if srgb is not None:
            base = srgb.attrib.get("val", "")
        else:
            sys_clr = node.find("a:sysClr", ns)
            base = sys_clr.attrib.get("lastClr", "") if sys_clr is not None else ""
        if not base:
            return ""
        return _apply_tint(base, tint)
    except Exception:
        return ""


def _apply_tint(hex_color: str, tint: float) -> str:
    hex_color = hex_color.replace("#", "")[-6:]
    if not hex_color or tint == 0:
        return hex_color.upper()
    channels = [int(hex_color[i:i + 2], 16) for i in (0, 2, 4)]
    adjusted = []
    for channel in channels:
        if tint < 0:
            value = int(channel * (1.0 + tint))
        else:
            value = int(channel * (1.0 - tint) + (255 - 255 * (1.0 - tint)))
        adjusted.append(max(0, min(255, value)))
    return "".join(f"{value:02X}" for value in adjusted)


def _specialty_from_header(header: str, fallback: str) -> str:
    fallback_name = re.sub(r"\.xlsx$", "", Path(str(fallback)).name, flags=re.I)
    fallback_lower = fallback_name.lower()
    filename_specialties = (
        ("generalist", "Generalist"),
        ("civil affairs", "Civil Affairs"),
        ("info analyst", "Information Analyst"),
        ("information analyst", "Information Analyst"),
        ("joint ops", "Joint Operations"),
        ("joint operations", "Joint Operations"),
        ("special assistant", "Special Assistant"),
        ("coordination", "Coordination"),
    )
    for marker, specialty in filename_specialties:
        if marker in fallback_lower:
            return specialty
    if not _is_usable_output_text(re.sub(r"(?i)responsibilities", "", header)):
        return "General"

    text = re.sub(r"\b(P\s*-?\s*[1-5]|D\s*-?\s*1|NO\s*-?\s*[A-D])\b", "", header, flags=re.I)
    text = re.sub(r"\bResponsibilities\b", "", text, flags=re.I)
    text = re.sub(r"\b(Associate|Senior|Assistant|Chief of Service)\b", "", text, flags=re.I)
    text = re.sub(r"\bOfficer\b", "", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip(" -_:")
    if _is_usable_output_text(text):
        return text

    fallback_name = re.sub(
        r"(?i)\b(transposing|ksa|mapping|theme|version|copy|output|6\.2|v\d+)\b",
        " ",
        fallback_name,
    )
    fallback_name = re.sub(r"[_-]+", " ", fallback_name)
    return re.sub(r"\s+", " ", fallback_name).strip() or "General"


def _canonical_theme_key(value: str) -> str:
    text = str(value or "").lower()
    text = re.sub(r"[^a-z0-9]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    aliases = {
        "information anaylst": "information analyst",
        "info analyst": "information analyst",
        "joint ops": "joint operations",
    }
    return aliases.get(text, text)


def _theme_color_map_from_workbook(wb) -> Dict[str, str]:
    colors: Dict[str, str] = {}
    for ws in wb.worksheets:
        if "theme" not in ws.title.lower():
            continue
        for row in range(1, ws.max_row + 1):
            name = _cell_text(ws, row, 2)
            fill = _cell_fill_hex(ws.cell(row, 2)) or _cell_fill_hex(ws.cell(row, 1))
            key = _canonical_theme_key(name)
            if key and fill:
                colors[key] = fill
    return colors


def _theme_fill_from_map(theme_name: str, theme_color_map: Dict[str, str]) -> str:
    key = _canonical_theme_key(theme_name)
    if key in theme_color_map:
        return theme_color_map[key]
    for known_key, fill in theme_color_map.items():
        if key and (key in known_key or known_key in key):
            return fill
    return ""


def _source_row_from_formula(formula: object) -> Tuple[str, Optional[int]]:
    text = str(formula or "")
    matches = re.findall(r"'?([^'=!]+)'?!\$?[A-Z]{1,3}\$?(\d+)", text)
    for sheet_name, row_text in reversed(matches):
        if sheet_name.lower().startswith("input_") and "ksa_check" in sheet_name.lower():
            return sheet_name, int(row_text)
    return "", None


def _find_row_with_labels(ws: Worksheet, required: Iterable[str]) -> Optional[int]:
    labels = tuple(x.lower() for x in required)
    for row in range(1, min(ws.max_row, 30) + 1):
        values = [
            _cell_text(ws, row, col).lower().replace("\n", " ")
            for col in range(1, min(ws.max_column, 12) + 1)
        ]
        if all(any(label in value for value in values) for label in labels):
            return row
    return None


def _is_usable_output_text(value: str) -> bool:
    text = str(value or "").strip()
    if not text or text.lower() in {"0", "[]", "[responsibilities]", "[ksa]"}:
        return False
    return not text.startswith("#")


def _parse_output_responsibilities(
    ws: Worksheet,
    source_name: str,
    theme_color_map: Optional[Dict[str, str]] = None,
    formula_ws: Optional[Worksheet] = None,
) -> Tuple[str, str, List[Responsibility]]:
    header_row = _find_row_with_labels(ws, ("area of work", "responsib"))
    if header_row is None:
        return "", "", []

    area_col = number_col = text_col = None
    for col in range(1, ws.max_column + 1):
        value = _cell_text(ws, header_row, col).lower()
        if "area of work" in value:
            area_col = col
        elif value in {"no.", "no", "nr", "nr2"}:
            number_col = col
        elif "responsib" in value:
            text_col = col
    if not all((area_col, number_col, text_col)):
        return "", "", []

    header = _cell_text(ws, header_row, text_col)
    specialty = _specialty_from_header(header, source_name)
    responsibilities: List[Responsibility] = []
    blank_run = 0
    current_area = ""
    current_area_fill = ""
    current_area_bold = False
    theme_color_map = theme_color_map or {}
    for row in range(header_row + 1, ws.max_row + 1):
        text = _cell_text(ws, row, text_col)
        if not _is_usable_output_text(text):
            blank_run += 1
            if blank_run >= 4:
                break
            continue
        blank_run = 0
        area = _cell_text(ws, row, area_col)
        area_cell = ws.cell(row, area_col)
        if area:
            current_area = area
            current_area_bold = bool(area_cell.font.bold)
        area_fill = _cell_fill_hex(area_cell)
        mapped_theme_fill = _theme_fill_from_map(current_area, theme_color_map)
        if area_fill:
            current_area_fill = area_fill
        elif mapped_theme_fill:
            current_area_fill = mapped_theme_fill
        raw_number = ws.cell(row, number_col).value
        try:
            number = int(float(raw_number))
        except (TypeError, ValueError):
            number = len(responsibilities) + 1
        text_cell = ws.cell(row, text_col)
        number_cell = ws.cell(row, number_col)
        source_mandatory = False
        if formula_ws is not None:
            source_sheet_name, source_row = _source_row_from_formula(formula_ws.cell(row, text_col).value)
            if source_sheet_name and source_row and source_sheet_name in ws.parent.sheetnames:
                source_ws = ws.parent[source_sheet_name]
                source_mandatory = str(source_ws.cell(source_row, 48).value or "").strip().lower() == "mandatory"
        theme_fill = area_fill or current_area_fill
        number_fill = _cell_fill_hex(number_cell) or theme_fill
        text_fill = _cell_fill_hex(text_cell)
        text_bold = bool(text_cell.font.bold) or source_mandatory
        responsibilities.append(
            Responsibility(
                number=number,
                theme=current_area or "General",
                text=text,
                mandatory=text_bold,
                fill_hex=theme_fill,
                theme_fill_hex=theme_fill,
                number_fill_hex=number_fill,
                text_fill_hex=text_fill,
                theme_bold=current_area_bold,
                number_bold=bool(number_cell.font.bold),
                text_bold=text_bold,
            )
        )
    return specialty, header, responsibilities


def _parse_output_ksa(ws: Worksheet) -> List[KsaRow]:
    header_row = _find_row_with_labels(ws, ("knowledge, skills", "responsib"))
    if header_row is None:
        return []

    number_col = text_col = mapping_start = None
    for col in range(1, ws.max_column + 1):
        value = _cell_text(ws, header_row, col).lower()
        if value in {"no.", "no", "nr", "nr2"}:
            number_col = col
        elif "knowledge, skills" in value:
            text_col = col
        elif "responsib" in value and mapping_start is None:
            mapping_start = col
    if not all((number_col, text_col, mapping_start)):
        return []

    rows: List[KsaRow] = []
    blank_run = 0
    for row in range(header_row + 1, ws.max_row + 1):
        text = _cell_text(ws, row, text_col)
        if not _is_usable_output_text(text):
            blank_run += 1
            if blank_run >= 4:
                break
            continue
        blank_run = 0
        raw_number = ws.cell(row, number_col).value
        try:
            number = int(float(raw_number))
        except (TypeError, ValueError):
            number = len(rows) + 1
        mapped: List[int] = []
        mapping_fills: Dict[int, str] = {}
        mapping_bolds: Dict[int, bool] = {}
        for col in range(mapping_start, ws.max_column + 1):
            mapping_cell = ws.cell(row, col)
            value = mapping_cell.value
            try:
                mapped_number = int(float(value))
            except (TypeError, ValueError):
                continue
            if mapped_number > 0 and mapped_number not in mapped:
                mapped.append(mapped_number)
                mapping_fills[mapped_number] = _cell_fill_hex(mapping_cell)
                mapping_bolds[mapped_number] = bool(mapping_cell.font.bold)
        text_cell = ws.cell(row, text_col)
        rows.append(
            KsaRow(
                number=number,
                text=text,
                responsibilities=mapped,
                text_bold=bool(text_cell.font.bold),
                text_fill_hex=_cell_fill_hex(text_cell),
                mapping_fills=mapping_fills,
                mapping_bolds=mapping_bolds,
            )
        )
    return rows


def parse_output_workbook(path: str | Path) -> Dict[str, List[LevelTables]]:
    """Read the cached values from Output_<level>_Resp and Output_<level>_KSA sheets."""
    wb = load_workbook(path, data_only=True)
    formula_wb = load_workbook(path, data_only=False)
    theme_color_map = _theme_color_map_from_workbook(wb)
    pairs: Dict[str, Dict[str, Worksheet]] = {}
    formula_pairs: Dict[str, Dict[str, Worksheet]] = {}
    for ws in wb.worksheets:
        match = OUTPUT_SHEET_RE.match(ws.title.strip())
        if not match:
            continue
        level = _normalize_level(match.group(1))
        kind = match.group(2).lower()
        pairs.setdefault(level, {})[kind] = ws
    for ws in formula_wb.worksheets:
        match = OUTPUT_SHEET_RE.match(ws.title.strip())
        if not match:
            continue
        level = _normalize_level(match.group(1))
        kind = match.group(2).lower()
        formula_pairs.setdefault(level, {})[kind] = ws

    result: Dict[str, List[LevelTables]] = {}
    for level, sheets in pairs.items():
        if "resp" not in sheets or "ksa" not in sheets:
            continue
        specialty, header_title, responsibilities = _parse_output_responsibilities(
            sheets["resp"],
            str(path),
            theme_color_map,
            formula_pairs.get(level, {}).get("resp"),
        )
        ksa_rows = _parse_output_ksa(sheets["ksa"])
        if responsibilities or ksa_rows:
            result.setdefault(level, []).append(
                LevelTables(
                    level=level,
                    specialty=specialty or _specialty_from_header("", str(path)),
                    responsibilities=responsibilities,
                    ksa_rows=ksa_rows,
                    header_title=header_title,
                )
            )
    return result


def parse_level_tables_from_upload(path: str | Path) -> Tuple[Dict[str, List[LevelTables]], List[str]]:
    """Accept one XLSX or a ZIP of XLSX files and combine all available specialties."""
    source = Path(path)
    workbook_paths: List[Path] = []
    temp_dir: Optional[Path] = None
    if source.suffix.lower() == ".zip":
        temp_dir = Path(tempfile.mkdtemp(prefix="gjp_excel_zip_"))
        with zipfile.ZipFile(source) as archive:
            for member in archive.infolist():
                member_path = Path(member.filename)
                if member.is_dir() or member_path.suffix.lower() != ".xlsx" or "__MACOSX" in member_path.parts:
                    continue
                safe_name = f"{len(workbook_paths):03d}_{member_path.name}"
                extracted = temp_dir / safe_name
                extracted.write_bytes(archive.read(member))
                workbook_paths.append(extracted)
    else:
        workbook_paths = [source]

    if not workbook_paths:
        raise ValueError("No .xlsx files were found in the uploaded file.")

    combined: Dict[str, List[LevelTables]] = {}
    warnings: List[str] = []
    for workbook_path in workbook_paths:
        parsed = parse_output_workbook(workbook_path)
        if not parsed:
            warnings.append(f"No usable Output_<level>_Resp/KSA sheet pairs found in {workbook_path.name}.")
            continue
        for level, tables in parsed.items():
            combined.setdefault(level, []).extend(tables)
    return combined, warnings


# -----------------------------------------------------------------------------
# Word table formatting helpers
# -----------------------------------------------------------------------------

def _set_cell_fill(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex.replace("#", ""))


def _set_cell_borders(cell, color="000000", size="6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def _format_cell_text(cell, font_size=8, bold=False, font_color: Optional[str] = None, align: Optional[int] = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.bold = bold
            if font_color:
                run.font.color.rgb = _rgb(font_color)


def _rgb(hex_color: str):
    from docx.shared import RGBColor
    h = hex_color.replace("#", "")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _block_element(block):
    """Return the underlying XML element for a Paragraph or Table."""
    if hasattr(block, "_p"):
        return block._p
    if hasattr(block, "_tbl"):
        return block._tbl
    raise TypeError(f"Unsupported block type: {type(block)}")


def _insert_paragraph_after(block, text: str = "", style=None):
    """
    Insert a paragraph immediately after an existing paragraph/table.

    This version is order-safe. The previous implementation always inserted after
    the same anchor, which could reverse the visual order when several blocks were
    inserted one after another.
    """
    parent = block._parent
    new_para = parent.add_paragraph()
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style

    _block_element(block).addnext(new_para._p)
    return new_para


def _insert_paragraph_before(paragraph, text: str = "", style=None):
    """Insert a paragraph immediately before an existing paragraph."""
    parent = paragraph._parent
    new_para = parent.add_paragraph()
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style

    paragraph._p.addprevious(new_para._p)
    return new_para


def _insert_table_after(block, rows: int, cols: int):
    """Insert a table immediately after an existing paragraph/table."""
    parent = block._parent
    table = parent.add_table(rows=rows, cols=cols, width=Inches(7.3))
    _block_element(block).addnext(table._tbl)
    return table


def _remove_element(element) -> None:
    element.getparent().remove(element)


def _paragraph_is_heading(paragraph, name: str) -> bool:
    return paragraph.text.strip().lower().startswith(name.lower())


def _clear_tables_after_heading_until(doc: DocumentObject, start_heading: str, stop_headings: Iterable[str]) -> Optional[object]:
    """
    Remove tables after a heading until a stop heading is reached.
    Returns the heading paragraph object.
    """
    body = doc._body._element
    children = list(body)
    start_idx = None
    for i, child in enumerate(children):
        if child.tag.endswith("}p"):
            p = None
            # Map XML paragraph to python-docx paragraph.
            for para in doc.paragraphs:
                if para._p is child:
                    p = para
                    break
            if p is not None and _paragraph_is_heading(p, start_heading):
                start_idx = i
                break
    if start_idx is None:
        return None

    # Find heading paragraph object.
    start_para = next((p for p in doc.paragraphs if p._p is children[start_idx]), None)
    stop_lower = tuple(s.lower() for s in stop_headings)
    for child in list(body)[start_idx + 1:]:
        if child.tag.endswith("}p"):
            p = next((para for para in doc.paragraphs if para._p is child), None)
            if p is not None and any(p.text.strip().lower().startswith(s) for s in stop_lower):
                break
        if child.tag.endswith("}tbl"):
            _remove_element(child)
    return start_para


def _find_heading_paragraph(doc: DocumentObject, heading: str) -> Optional[object]:
    for paragraph in doc.paragraphs:
        if _paragraph_is_heading(paragraph, heading):
            return paragraph
    return None


def _replace_section_paragraph_text(doc: DocumentObject, heading: str, replacement_text: str, stop_headings: Iterable[str]) -> None:
    """
    Replace the text immediately under a section heading while keeping the
    heading in its original template position.

    This keeps the template order:
        Work Interactions -> Introduction -> Responsibilities -> Skills

    It removes old paragraphs, including blank spacer paragraphs, until the next
    stop heading/table, then inserts the new paragraph directly after the heading.
    """
    heading_p = _find_heading_paragraph(doc, heading)
    if heading_p is None or not replacement_text:
        return

    body = doc._body._element
    children = list(body)
    start_idx = children.index(heading_p._p)
    stop_lower = tuple(s.lower() for s in stop_headings)

    to_remove = []

    for child in children[start_idx + 1:]:
        if child.tag.endswith("}p"):
            p = next((para for para in doc.paragraphs if para._p is child), None)
            if p is not None:
                p_text = p.text.strip()
                # Stop at the next major section heading.
                if any(p_text.lower().startswith(s) for s in stop_lower):
                    break
                # Remove both old body text and blank spacing paragraphs.
                to_remove.append(child)

        elif child.tag.endswith("}tbl"):
            # Stop before the next table. Responsibilities/Skills tables are handled separately.
            break

    for child in to_remove:
        _remove_element(child)

    new_p = _insert_paragraph_after(heading_p, replacement_text)
    for run in new_p.runs:
        run.font.size = Pt(11)


# -----------------------------------------------------------------------------
# Build and insert Responsibility/KSA tables
# -----------------------------------------------------------------------------

def add_responsibility_table_after(
    paragraph,
    responsibilities: List[Responsibility],
    header_title: str = "[Responsibilities]",
) -> object:
    table = _insert_table_after(paragraph, rows=len(responsibilities) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [1.6, 0.5, 5.2]

    headers = ["Area of Work", "No.", header_title]
    for c, header in enumerate(headers):
        cell = table.cell(0, c)
        cell.text = header
        _set_cell_fill(cell, HEADER_FILL)
        _set_cell_borders(cell)
        _set_cell_width(cell, widths[c])
        _format_cell_text(cell, font_size=8, bold=True, font_color=HEADER_FONT, align=WD_ALIGN_PARAGRAPH.CENTER)

    for r, resp in enumerate(responsibilities, start=1):
        row_values = [resp.theme, str(resp.number), resp.text]
        fills = [
            resp.theme_fill_hex or resp.fill_hex or _theme_color(resp.theme),
            resp.number_fill_hex or resp.fill_hex or _theme_color(resp.theme),
            resp.text_fill_hex,
        ]
        bolds = [resp.theme_bold, resp.number_bold, resp.text_bold]
        for c, value in enumerate(row_values):
            cell = table.cell(r, c)
            cell.text = value
            _set_cell_borders(cell)
            _set_cell_width(cell, widths[c])
            if fills[c]:
                _set_cell_fill(cell, fills[c])
            _format_cell_text(
                cell,
                font_size=8,
                bold=bool(bolds[c]),
                align=WD_ALIGN_PARAGRAPH.CENTER if c in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT,
            )
    return table


def _chunks(items: List[int], size: int) -> List[List[int]]:
    return [items[i:i + size] for i in range(0, len(items), size)]


def add_ksa_table_after(
    paragraph,
    ksa_rows: List[KsaRow],
    responsibilities: List[Responsibility],
    header_title: str = "Responsibilities",
    max_numbers_per_line: Optional[int] = None,
) -> object:
    """
    Build the Skills/KSA mapping table.

    Important:
    - Keep ALL mapped responsibility numbers.
    - If one KSA maps to more than max_numbers_per_line responsibilities,
      add continuation rows instead of deleting the extra numbers.
    - The generated GJP table omits the Excel No. column.
      Continuation rows keep the KSA text blank and continue the remaining numbers.
    """
    resp_by_no = {r.number: r for r in responsibilities}
    if max_numbers_per_line is None:
        max_numbers_per_line = max(
            1,
            max((len(row.responsibilities) for row in ksa_rows), default=1),
        )

    # Build expanded rows first so we know the total row count.
    expanded_rows = []
    for ksa in ksa_rows:
        nums = list(ksa.responsibilities or [])
        chunks = _chunks(nums, max_numbers_per_line) if nums else [[]]
        for chunk_index, chunk in enumerate(chunks):
            expanded_rows.append((ksa, chunk, chunk_index))

    # 1 fixed KSA column + responsibility-number columns. The Excel No. column
    # is intentionally not copied into the generated GJP document.
    total_cols = 1 + max_numbers_per_line
    table = _insert_table_after(paragraph, rows=len(expanded_rows) + 1, cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    number_width = max(0.22, 4.35 / max_numbers_per_line)
    widths = [2.95] + [number_width] * max_numbers_per_line

    # Header
    header_cells = ["Knowledge, Skills, and Abilities", header_title] + [""] * (max_numbers_per_line - 1)
    for c, header in enumerate(header_cells):
        cell = table.cell(0, c)
        cell.text = header
        _set_cell_fill(cell, HEADER_FILL)
        _set_cell_borders(cell)
        _set_cell_width(cell, widths[c])
        _format_cell_text(
            cell,
            font_size=8,
            bold=True,
            font_color=HEADER_FONT,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    # Body
    for r, (ksa, chunk, chunk_index) in enumerate(expanded_rows, start=1):
        # Only the first visual row displays KSA text.
        ksa_cell = table.cell(r, 0)
        ksa_cell.text = ksa.text if chunk_index == 0 else ""
        _set_cell_borders(ksa_cell)
        _set_cell_width(ksa_cell, widths[0])
        if ksa.text_fill_hex:
            _set_cell_fill(ksa_cell, ksa.text_fill_hex)
        _format_cell_text(
            ksa_cell,
            font_size=7,
            bold=bool(ksa.text_bold),
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        # Fill responsibility numbers for this chunk.
        for offset in range(max_numbers_per_line):
            c = 1 + offset
            cell = table.cell(r, c)
            mapped_resp = None

            if offset < len(chunk):
                n = chunk[offset]
                cell.text = str(n)
                mapped_resp = resp_by_no.get(n)
                _set_cell_fill(
                    cell,
                    ksa.mapping_fills.get(n)
                    or (mapped_resp.fill_hex if mapped_resp else "")
                    or _theme_color(mapped_resp.theme if mapped_resp else ""),
                )
            else:
                cell.text = ""

            _set_cell_borders(cell)
            _set_cell_width(cell, widths[c])
            _format_cell_text(
                cell,
                font_size=7,
                bold=bool(ksa.mapping_bolds.get(chunk[offset], False) if offset < len(chunk) else False)
                or bool(mapped_resp and mapped_resp.mandatory),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

    if max_numbers_per_line > 1:
        merged = table.cell(0, 1).merge(table.cell(0, total_cols - 1))
        merged.text = header_title
        _set_cell_fill(merged, HEADER_FILL)
        _set_cell_borders(merged)
        _format_cell_text(
            merged,
            font_size=8,
            bold=True,
            font_color=HEADER_FONT,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    return table


def _is_general_specialty_name(name: str) -> bool:
    return not str(name or "").strip() or str(name).strip().lower() in {"general", "generalist"}


def _format_section_heading(paragraph) -> None:
    if paragraph.runs:
        paragraph.runs[0].bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _canonical_specialty(value: str) -> str:
    text = str(value or "").lower()
    text = re.sub(r"\bspecialt(?:y|ies)\b", " ", text)
    text = re.sub(r"\b(associate|assistant|senior|chief|service|officer|the)\b", " ", text)
    text = re.sub(r"[^a-z0-9]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    aliases = {
        "general": "generalist",
        "generalist": "generalist",
        "info analyst": "information analyst",
        "information analysis": "information analyst",
        "joint ops": "joint operations",
        "coordination officer": "coordination",
    }
    return aliases.get(text, text)


def _specialty_match_score(template_name: str, uploaded_name: str) -> int:
    template_key = _canonical_specialty(template_name)
    uploaded_key = _canonical_specialty(uploaded_name)
    if not template_key or not uploaded_key:
        return 0
    if template_key == uploaded_key:
        return 100
    if template_key in uploaded_key or uploaded_key in template_key:
        return 70
    template_words = set(template_key.split())
    uploaded_words = set(uploaded_key.split())
    return len(template_words & uploaded_words) * 10


def _introduction_for_specialty(
    introductions: Optional[Dict[str, Dict[str, str]]],
    level: str,
    specialty: str,
) -> str:
    if not introductions:
        return ""
    by_specialty = introductions.get(level) or {}
    if not by_specialty:
        equivalent = LEVEL_EQUIVALENT.get(level)
        by_specialty = introductions.get(equivalent or "") or {}
    if not by_specialty:
        return ""

    best_text = ""
    best_score = 0
    for intro_specialty, intro_text in by_specialty.items():
        if intro_specialty.lower().startswith("all specialt"):
            continue
        score = _specialty_match_score(specialty, intro_specialty)
        if score > best_score:
            best_score = score
            best_text = intro_text
    if best_score >= 10:
        return best_text
    return by_specialty.get("") or by_specialty.get("Introduction", "")


def _replace_intro_inside_segment(
    body,
    segment,
    paragraph_by_xml,
    introduction_text: str,
    stop_texts: Iterable[str] = ("responsibilities",),
) -> None:
    if not introduction_text:
        return

    intro_heading = None
    intro_index = None
    for index, child in enumerate(segment):
        paragraph = paragraph_by_xml.get(child)
        if paragraph is not None and paragraph.text.strip().lower() == "introduction":
            intro_heading = paragraph
            intro_index = index
            break
    if intro_heading is None or intro_index is None:
        return

    stop_lower = tuple(s.lower() for s in stop_texts)
    to_remove = []
    for child in segment[intro_index + 1:]:
        paragraph = paragraph_by_xml.get(child)
        if paragraph is not None and any(paragraph.text.strip().lower().startswith(s) for s in stop_lower):
            break
        if child.tag.endswith("}p"):
            to_remove.append(child)
        elif child.tag.endswith("}tbl"):
            break

    for child in to_remove:
        if child.getparent() is body:
            _remove_element(child)
    new_p = _insert_paragraph_after(intro_heading, introduction_text)
    for run in new_p.runs:
        run.font.size = Pt(11)


def _specialty_blocks(doc: DocumentObject) -> List[Tuple[str, object, int, int]]:
    body = doc._body._element
    children = list(body)
    paragraph_by_xml = {p._p: p for p in doc.paragraphs}
    starts: List[Tuple[str, object, int]] = []
    for index, child in enumerate(children):
        paragraph = paragraph_by_xml.get(child)
        if paragraph is None:
            continue
        match = re.match(r"\s*Specialt(?:y|ies)\s*:\s*(.+)", paragraph.text, flags=re.I)
        if match:
            starts.append((match.group(1).strip(), paragraph, index))
    blocks: List[Tuple[str, object, int, int]] = []
    for index, (name, paragraph, start) in enumerate(starts):
        end = starts[index + 1][2] if index + 1 < len(starts) else len(children)
        blocks.append((name, paragraph, start, end))
    return blocks


def _replace_specialty_blocks(
    doc: DocumentObject,
    level_tables: List[LevelTables],
    target_level: str,
    gjp_area: str,
    introductions: Optional[Dict[str, Dict[str, str]]] = None,
) -> None:
    blocks = _specialty_blocks(doc)
    if not blocks:
        replace_tables_in_document(doc, level_tables, target_level, gjp_area, introductions)
        return

    remaining = list(level_tables)
    assignments: List[Tuple[Tuple[str, object, int, int], Optional[LevelTables]]] = []
    for block in blocks:
        best = None
        best_score = 0
        for table_data in remaining:
            score = _specialty_match_score(block[0], table_data.specialty)
            if score > best_score:
                best = table_data
                best_score = score
        if best is not None and best_score >= 10:
            remaining.remove(best)
            assignments.append((block, best))
        else:
            assignments.append((block, None))

    if len(level_tables) == 1 and all(table_data is None for _, table_data in assignments):
        # A workbook without specialties still belongs to one department/workstream.
        # Put its tables in the first template specialty block and remove the rest.
        assignments[0] = (blocks[0], level_tables[0])
        first_specialty_paragraph = blocks[0][1]
        if _is_general_specialty_name(level_tables[0].specialty):
            first_specialty_paragraph.text = f"Specialty: {gjp_area}"
            _format_section_heading(first_specialty_paragraph)

    body = doc._body._element
    original_children = list(body)
    paragraph_by_xml = {p._p: p for p in doc.paragraphs}

    # Delete template specialty blocks that have no uploaded data for this level.
    for (name, paragraph, start, end), table_data in reversed(assignments):
        if table_data is None:
            for child in original_children[start:end]:
                if child.getparent() is body:
                    _remove_element(child)

    # Replace the two tables inside each retained specialty block.
    last_block = None
    for (name, specialty_paragraph, start, end), table_data in assignments:
        if table_data is None or specialty_paragraph._p.getparent() is not body:
            continue
        segment = [child for child in original_children[start:end] if child.getparent() is body]
        resp_heading = None
        ksa_heading = None
        for child in segment:
            paragraph = paragraph_by_xml.get(child)
            if paragraph is None:
                continue
            text = paragraph.text.strip().lower()
            if text == "responsibilities":
                resp_heading = paragraph
            elif text == "skills" or text.startswith("knowledge, skills"):
                ksa_heading = paragraph
        intro_text = _introduction_for_specialty(introductions, target_level, table_data.specialty)
        _replace_intro_inside_segment(body, segment, paragraph_by_xml, intro_text)
        for child in segment:
            if child.tag.endswith("}tbl"):
                _remove_element(child)
        if resp_heading is None or ksa_heading is None:
            continue

        header_title = table_data.header_title or f"{target_level} {gjp_area} Responsibilities"
        source_level = detect_template_level(doc)
        header_title = re.sub(
            rf"\b{re.escape(source_level)}\b",
            target_level,
            header_title,
            flags=re.I,
        )
        add_responsibility_table_after(resp_heading, table_data.responsibilities, header_title)
        last_block = add_ksa_table_after(ksa_heading, table_data.ksa_rows, table_data.responsibilities, header_title)

    if last_block is None:
        retained_blocks = [
            original_children[end - 1]
            for (_, paragraph, start, end), table_data in assignments
            if table_data is not None and paragraph._p.getparent() is body and end > start
        ]
        if retained_blocks:
            last_block = retained_blocks[-1]

    for table_data in remaining:
        if last_block is None:
            last_block = doc.add_paragraph()

        spec_p = _insert_paragraph_after(last_block, f"Specialty: {table_data.specialty}")
        _format_section_heading(spec_p)
        last_block = spec_p

        intro_text = _introduction_for_specialty(introductions, target_level, table_data.specialty)
        if intro_text:
            intro_p = _insert_paragraph_after(last_block, "Introduction")
            _format_section_heading(intro_p)
            last_block = _insert_paragraph_after(intro_p, intro_text)

        resp_p = _insert_paragraph_after(last_block, "Responsibilities")
        _format_section_heading(resp_p)
        header_title = table_data.header_title or f"{target_level} {gjp_area} Responsibilities"
        last_block = add_responsibility_table_after(resp_p, table_data.responsibilities, header_title)

        skill_p = _insert_paragraph_after(last_block, "Knowledge, Skills, Abilities (KSAs)")
        _format_section_heading(skill_p)
        last_block = add_ksa_table_after(skill_p, table_data.ksa_rows, table_data.responsibilities, header_title)


def _remove_all_responsibility_ksa_sections(doc: DocumentObject) -> None:
    """Keep each specialty/introduction, but remove Responsibility and KSA sections."""
    body = doc._body._element
    children = list(body)
    paragraph_by_xml = {p._p: p for p in doc.paragraphs}
    blocks = _specialty_blocks(doc)

    if blocks:
        for _, _, start, end in reversed(blocks):
            remove_from = None
            for index in range(start, end):
                paragraph = paragraph_by_xml.get(children[index])
                if paragraph is not None and paragraph.text.strip().lower() == "responsibilities":
                    remove_from = index
                    break
            if remove_from is not None:
                for child in children[remove_from:end]:
                    if child.getparent() is body:
                        _remove_element(child)
        return

    # Templates without specialty labels: remove from Responsibilities through
    # the end of the KSA section, stopping at the next major section if present.
    start = None
    end = len(children)
    for index, child in enumerate(children):
        paragraph = paragraph_by_xml.get(child)
        if paragraph is None:
            continue
        text = paragraph.text.strip().lower()
        if start is None and text == "responsibilities":
            start = index
            continue
        if start is not None and text in {"education", "work experience", "work interactions"}:
            end = index
            break
    if start is not None:
        for child in children[start:end]:
            if child.getparent() is body:
                _remove_element(child)


def replace_tables_in_document(
    doc: DocumentObject,
    level_tables: List[LevelTables],
    target_level: str = "",
    gjp_area: str = "",
    introductions: Optional[Dict[str, Dict[str, str]]] = None,
) -> None:
    """
    Replace the template Responsibilities and Skills tables without changing the
    template section order.

    The generated document follows the template order exactly:

        Education
        Work Experience
        Work Interactions
        Introduction
        Responsibilities
        Skills

    For the first specialty, the existing template Responsibilities and Skills
    headings are reused in place. Extra specialty blocks, if any, are appended
    after the first Skills table in the same internal order:
        Specialty -> Responsibilities -> Skills.
    """
    if not level_tables:
        return

    # Keep the original headings exactly where they are in the template.
    # Only remove the old tables below those headings.
    resp_heading = _clear_tables_after_heading_until(doc, "Responsibilities", ["Skills"])
    skills_heading = _clear_tables_after_heading_until(doc, "Skills", ["**Note", "Note", "Education", "Work Experience", "Work Interactions", "Introduction", "Responsibilities"])

    if resp_heading is None:
        resp_heading = doc.add_paragraph("Responsibilities")
        _format_section_heading(resp_heading)
    if skills_heading is None:
        skills_heading = doc.add_paragraph("Skills")
        _format_section_heading(skills_heading)

    # First specialty uses the original template locations.
    first = level_tables[0]
    first_intro = _introduction_for_specialty(introductions, target_level, first.specialty)
    if first_intro:
        _replace_section_paragraph_text(doc, "Introduction", first_intro, ["Responsibilities", "Skills"])

    # If the first table has a real specialty name, place it immediately before
    # Responsibilities, so the order remains:
    # Introduction -> Specialty -> Responsibilities -> Skills.
    if not _is_general_specialty_name(first.specialty):
        spec_p = _insert_paragraph_before(resp_heading, first.specialty)
        _format_section_heading(spec_p)

    first_header = first.header_title or f"{target_level} {gjp_area} Responsibilities".strip()
    add_responsibility_table_after(resp_heading, first.responsibilities, first_header)
    last_block = add_ksa_table_after(skills_heading, first.ksa_rows, first.responsibilities, first_header)

    # Additional specialties are appended after the previous Skills table.
    # Each appended block is inserted sequentially after the previous block
    # to avoid reversed order.
    for tbl in level_tables[1:]:
        if not _is_general_specialty_name(tbl.specialty):
            spec_p = _insert_paragraph_after(last_block, tbl.specialty)
            _format_section_heading(spec_p)
            last_block = spec_p

        intro_text = _introduction_for_specialty(introductions, target_level, tbl.specialty)
        if intro_text:
            intro_p = _insert_paragraph_after(last_block, "Introduction")
            _format_section_heading(intro_p)
            last_block = _insert_paragraph_after(intro_p, intro_text)

        resp_p = _insert_paragraph_after(last_block, "Responsibilities")
        _format_section_heading(resp_p)
        table_header = tbl.header_title or f"{target_level} {gjp_area} Responsibilities".strip()
        last_block = add_responsibility_table_after(resp_p, tbl.responsibilities, table_header)

        skill_p = _insert_paragraph_after(last_block, "Skills")
        _format_section_heading(skill_p)
        last_block = add_ksa_table_after(skill_p, tbl.ksa_rows, tbl.responsibilities, table_header)


# -----------------------------------------------------------------------------
# Main generation workflow
# -----------------------------------------------------------------------------

def generate_one_document(
    template_path: str | Path,
    target_level: str,
    level_tables: List[LevelTables],
    gjp_area: str,
    job_info: Optional[JobInfo] = None,
    source_job_info: Optional[JobInfo] = None,
    interactions: Optional[Dict[str, str]] = None,
    introductions: Optional[Dict[str, Dict[str, str]]] = None,
) -> bytes:
    doc = Document(str(template_path))
    source_level = detect_template_level(doc)

    replace_level_everywhere(doc, source_level, target_level)
    _replace_job_info(doc, source_level, target_level, job_info, source_job_info)
    replace_work_experience_years(doc, target_level)

    # Use equivalent P-level text for NO-levels if exact NO text is not provided.
    lookup_level = target_level
    equivalent = LEVEL_EQUIVALENT.get(target_level)
    interaction_text = (interactions or {}).get(lookup_level) or ((interactions or {}).get(equivalent) if equivalent else None)

    if interaction_text:
        _replace_section_paragraph_text(
            doc,
            "Work Interactions",
            interaction_text,
            ["Education", "Work Experience", "Specialty", "Introduction", "Responsibilities", "Skills"],
        )

    if level_tables:
        _replace_specialty_blocks(doc, level_tables, target_level, gjp_area, introductions)
    else:
        _remove_all_responsibility_ksa_sections(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _safe_filename_part(value: str) -> str:
    value = re.sub(r'[<>:"/\\|?*]+', " ", str(value))
    return re.sub(r"\s+", " ", value).strip(" .") or "GJP"


def generate_documents(
    job_info_path: str | Path,
    excel_or_zip_path: str | Path,
    template_path: str | Path,
    intro_interaction_path: str | Path,
) -> Tuple[Dict[str, bytes], List[str]]:
    job_infos = parse_job_info_workbook(job_info_path)
    all_tables, warnings = parse_level_tables_from_upload(excel_or_zip_path)
    interactions, introductions = parse_intro_interaction_docx(intro_interaction_path)
    template_level = detect_template_level(Document(str(template_path)))
    source_job_info = job_infos.get(template_level)

    documents: Dict[str, bytes] = {}
    for level in OUTPUT_LEVELS:
        tables = all_tables.get(level, [])
        if not tables:
            continue
        job_info = job_infos.get(level)
        gjp_area = _strip_level_from_title(job_info.job_title, level) if job_info else "GJP"
        doc_bytes = generate_one_document(
            template_path=template_path,
            target_level=level,
            level_tables=tables,
            gjp_area=gjp_area,
            job_info=job_info,
            source_job_info=source_job_info,
            interactions=interactions,
            introductions=introductions,
        )
        if job_info and job_info.file_name:
            file_name = _safe_filename_part(Path(job_info.file_name).stem) + ".docx"
        else:
            file_name = f"{_safe_filename_part(gjp_area)} GJP 2.0 - {level}.docx"
        documents[file_name] = doc_bytes

    if not all_tables:
        raise ValueError(
            "No populated Output_<level>_Resp and Output_<level>_KSA sheet pair "
            "was found in the uploaded Excel file(s)."
        )
    if not documents:
        raise ValueError("No documents were generated because none of the uploaded table levels matched the configured output levels.")
    return documents, warnings


def build_documents_zip(documents: Dict[str, bytes]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for file_name, doc_bytes in documents.items():
            archive.writestr(file_name, doc_bytes)
    buffer.seek(0)
    return buffer.getvalue()


# -----------------------------------------------------------------------------
# Streamlit UI
# -----------------------------------------------------------------------------

def _save_uploaded_file(uploaded_file, suffix: str) -> Path:
    tmp_dir = Path(tempfile.mkdtemp(prefix="gjp_doc_generator_"))
    path = tmp_dir / f"uploaded{suffix}"
    path.write_bytes(uploaded_file.getbuffer())
    return path


def render_gjp_document_generator() -> None:
    if st is None:
        raise RuntimeError("Streamlit is not installed. Run: pip install streamlit")
    st.header("Generate GJP Word Documents")
    st.write(
        "Upload job title/job code information, specialty mapping Excel files, "
        "the introduction/work interactions Word file, and one GJP Word template. "
        "The tool generates a ZIP containing only the levels that have populated "
        "Responsibility/KSA output tables."
    )
    job_info_file = st.file_uploader(
        "1) Upload job title/job code file (.xlsx)",
        type=["xlsx"],
        key="gjp_job_info_xlsx",
    )
    excel_file = st.file_uploader(
        "2) Upload specialty Excel ZIP or one Excel file (.zip or .xlsx)",
        type=["zip", "xlsx"],
        key="gjp_level_excel_or_zip",
    )
    intro_file = st.file_uploader(
        "3) Upload introduction and interactions (.docx)",
        type=["docx"],
        key="intro_interaction_docx",
    )
    template_file = st.file_uploader(
        "4) Upload GJP Word template (.docx)",
        type=["docx"],
        key="gjp_template_docx",
    )

    if st.button("Generate GJP Documents", type="primary"):
        if job_info_file is None or excel_file is None or template_file is None or intro_file is None:
            st.error(
                "Please upload the job info Excel, specialty Excel/ZIP file, "
                "Work Interactions/Introduction Word file, and Word template."
            )
            return
        try:
            with st.spinner("Generating Word documents..."):
                job_info_path = _save_uploaded_file(job_info_file, ".xlsx")
                input_suffix = Path(excel_file.name).suffix.lower()
                excel_path = _save_uploaded_file(excel_file, input_suffix)
                template_path = _save_uploaded_file(template_file, ".docx")
                intro_path = _save_uploaded_file(intro_file, ".docx")
                documents, warnings = generate_documents(
                    job_info_path,
                    excel_path,
                    template_path,
                    intro_path,
                )

            if warnings:
                with st.expander("Generation warnings"):
                    for warning in warnings:
                        st.warning(warning)
            zip_bytes = build_documents_zip(documents)
            st.success(f"Done! Generated {len(documents)} Word document(s).")
            st.download_button(
                label="Download generated GJP documents ZIP",
                data=zip_bytes,
                file_name="generated_GJP_documents.zip",
                mime="application/zip",
                key="download_generated_gjp_documents_zip",
            )
        except Exception as exc:
            st.exception(exc)


if __name__ == "__main__":
    render_gjp_document_generator()
