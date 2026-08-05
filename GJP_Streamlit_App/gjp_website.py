
import re
import io
import zipfile
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from openpyxl import Workbook, load_workbook
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from openpyxl.styles import Font, Border, Side, Alignment, PatternFill
from openpyxl.utils import get_column_letter


# =========================
# Helper functions
# =========================

def clean_text(text):
    """Clean repeated spaces and line breaks."""
    if text is None:
        return ""
    text = str(text).replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def normalize_for_unique(text):
    """
    Normalize text for duplicate checking.
    Treat punctuation-only differences as the same.
    Example:
        "Conduct research, analysis."
        "Conduct research analysis"
    will be treated as similar/same for exact normalized matching.
    """
    text = clean_text(text).lower()

    # Replace punctuation with space
    text = re.sub(r"[^\w\s]", " ", text)

    # Normalize spaces
    text = re.sub(r"\s+", " ", text).strip()

    return text



def get_docx_text_and_tables(file_bytes):
    """
    Read all paragraphs and table cell texts from a docx file.
    Returns:
        paragraphs: list of paragraph texts
        all_text: combined text
        tables_text: list of rows, each row is list of cell texts
    """
    doc = Document(io.BytesIO(file_bytes))

    paragraphs = []
    for p in doc.paragraphs:
        txt = clean_text(p.text)
        if txt:
            paragraphs.append(txt)

    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_values = [clean_text(cell.text) for cell in row.cells]
            if any(row_values):
                tables_text.append(row_values)
                # Also add table cells into paragraph-like text for easier searching
                for cell_text in row_values:
                    if cell_text:
                        paragraphs.append(cell_text)

    all_text = "\n".join(paragraphs)
    return paragraphs, all_text, tables_text


def extract_gjp_job_title(paragraphs):
    """
    Extract GJP Job Title.

    Priority 1:
        From description sentence:
        The Chief of Service, Human Rights D-1 GJP is available for ...
        -> Chief of Service, Human Rights D-1

    Priority 2:
        From line under GENERIC JOB PROFILE:
        GENERIC JOB PROFILE
        Assistant Human Rights Officer: P-1

    Output will also normalize:
        Assistant Human Rights Officer: P-1
        -> Assistant Human Rights Officer P-1
    """

    # Priority 1: extract from "The ... GJP is available for ..."
    for txt in paragraphs:
        text = clean_text(txt)

        match = re.search(
            r"^The\s+(.+?)\s+GJP\s+is\s+available\s+for\b",
            text,
            re.IGNORECASE
        )

        if match:
            title = clean_text(match.group(1))
            title = title.replace(": ", " ")
            title = title.replace(" :", " ")
            title = re.sub(r"\s+", " ", title).strip()
            return title

    # Priority 2: extract the second line under GENERIC JOB PROFILE
    for i, txt in enumerate(paragraphs):
        if clean_text(txt).upper() == "GENERIC JOB PROFILE":
            for j in range(i + 1, min(i + 6, len(paragraphs))):
                candidate = clean_text(paragraphs[j])

                if candidate and not candidate.lower().startswith("version"):
                    candidate = candidate.replace(": ", " ")
                    candidate = candidate.replace(" :", " ")
                    candidate = re.sub(r"\s+", " ", candidate).strip()
                    return candidate

    # Priority 3 fallback: find something like "Officer: P-1" or "Officer P-1"
    pattern = re.compile(r".+\s*:?\s*(P|D|G|NO|FS)-?\d+", re.IGNORECASE)

    for txt in paragraphs:
        text = clean_text(txt)
        if pattern.match(text):
            text = text.replace(": ", " ")
            text = text.replace(" :", " ")
            text = re.sub(r"\s+", " ", text).strip()
            return text

    return ""


def normalize_workstream_name(folder_name):
    """
    Convert GJP Area/ Workstream into lowercase underscore format.
    Example:
        Human Rights -> human_rights
        Political Affairs -> political_affairs
    """
    text = clean_text(folder_name)
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    text = re.sub(r"_+", "_", text)
    text = text.strip("_")
    return text


def extract_title_before_level_from_description(paragraphs):
    """
    Extract title before level from:
        The Chief of Service, Human Rights D-1 GJP is available for ...
    Output:
        Chief of Service, Human Rights

    Supported levels:
        P-1, P-2, P-3, P-4, P-5,
        D-1,
        NO-A, NO-B, NO-C, NO-D
    """

    level_pattern = r"(?:P-[1-5]|D-1|NO-[A-D])"

    for txt in paragraphs:
        text = clean_text(txt)

        match = re.search(
            rf"^The\s+(.+?)\s+{level_pattern}\s+GJP\s+is\s+available\s+for\b",
            text,
            re.IGNORECASE
        )

        if match:
            return clean_text(match.group(1))

    return ""


def format_department_name(folder_name):
    """
    Convert folder / department name into required display format.
    Example:
        Human Rights -> Human rights
        Political Affairs -> Political affairs
    """
    text = clean_text(folder_name)
    text = re.sub(r"\s+", " ", text).strip()

    if not text:
        return ""

    # Make first letter uppercase, rest lowercase
    return text[0].upper() + text[1:].lower()


def clean_mapped_specialty_name(specialty, folder_name, paragraphs):
    """
    Rules:
    1. If specialty is General / general / Generalist:
       -> Generalist - [department name]
       Example:
          Human Rights -> Generalist - Human rights

    2. If specialty is blank:
       -> extract title before level from description sentence.
       Example:
          The Chief of Service, Human Rights D-1 GJP is available for
          -> Chief of Service, Human Rights
    """

    specialty = clean_text(specialty)
    specialty = specialty.replace('"', "").replace("“", "").replace("”", "").strip()

    # Rule 1: General / Generalist
    if specialty.lower() in ["general", "generalist"]:
        department_name = format_department_name(folder_name)
        return f"Generalist - {department_name}"

    # Rule 2: blank mapped specialty
    if specialty == "":
        extracted = extract_title_before_level_from_description(paragraphs)
        if extracted:
            return extracted

    return specialty


def extract_job_code(all_text):
    """
    Extract one or multiple Job Codes.
    Examples:
        The CCOG Code is 1.G.02. and the Job Code is 1220.
        The Job Codes are 10225, 3844 and 8624.
    Output:
        1220
        10225, 3844, 8624
    """
    text = clean_text(all_text)

    # Match "Job Code is ..." or "Job Codes are ..."
    match = re.search(
        r"Job Codes?\s+(?:is|are)\s+(.+?)(?:\.|\n|Signed:|The\s)",
        text,
        re.IGNORECASE
    )

    if not match:
        return ""

    segment = match.group(1)

    # Extract all number-like job codes
    codes = re.findall(r"\b\d+\b", segment)

    if codes:
        return ", ".join(codes)

    # fallback
    segment = segment.replace(" and ", ", ")
    segment = re.sub(r"\s+", " ", segment)
    return segment.strip().rstrip(".")


def extract_description(paragraphs):
    """
    Extract the full sentence/paragraph:
    The ... is available for ...
    """
    for txt in paragraphs:
        if re.search(r"^The\s+.+?\s+is available for\s+", txt, re.IGNORECASE):
            return clean_text(txt)

    # fallback: paragraph containing "available for use"
    for txt in paragraphs:
        if "available for use" in txt.lower():
            return clean_text(txt)

    return ""


def extract_section_paragraphs(paragraphs, section_name, stop_sections):
    """
    Extract paragraphs between one heading and the next heading.
    Example:
        section_name = "Education"
        stop_sections = ["Work Experience", "Work Interactions", "Introduction"]
    """
    start_idx = None

    for i, txt in enumerate(paragraphs):
        if clean_text(txt).lower() == section_name.lower():
            start_idx = i
            break

    if start_idx is None:
        return []

    results = []
    stop_lower = [s.lower() for s in stop_sections]

    for txt in paragraphs[start_idx + 1:]:
        current = clean_text(txt)
        if current.lower() in stop_lower:
            break
        if current:
            results.append(current)

    return results


def extract_education_description(paragraphs):
    """
    Education part:
    Usually first paragraph is:
      The education requirement is the same...
    Second paragraph is:
      Advanced university degree ...
    We need the whole second paragraph.
    """
    edu_paras = extract_section_paragraphs(
        paragraphs,
        "Education",
        ["Work Experience", "Work Interactions", "Introduction", "Responsibilities",
         "Knowledge, Skills, Abilities (KSAs)"]
    )

    if len(edu_paras) >= 2:
        return clean_text(edu_paras[1])
    elif len(edu_paras) == 1:
        return clean_text(edu_paras[0])
    return ""


def extract_minimum_education_level(education_description):
    """
    From:
      Advanced university degree (Master’s degree or equivalent) in human rights, law, ...
    Extract:
      Advanced university degree (Master’s degree or equivalent)
    """
    text = clean_text(education_description)

    # most common split: before " in ..."
    match = re.match(r"(.+?\))\s+in\s+", text, re.IGNORECASE)
    if match:
        return clean_text(match.group(1))

    # fallback: before first " in "
    parts = re.split(r"\s+in\s+", text, maxsplit=1, flags=re.IGNORECASE)
    if parts:
        return clean_text(parts[0])

    return text


def extract_work_experience_description(paragraphs):
    """
    Work Experience part:
    Usually first paragraph is:
      The work experience requirement is the same...
    Second paragraph is the actual requirement.
    We need the whole second paragraph.
    """
    work_paras = extract_section_paragraphs(
        paragraphs,
        "Work Experience",
        ["Work Interactions", "Introduction", "Responsibilities", "Knowledge, Skills, Abilities (KSAs)", "Education"]
    )

    # remove guidance paragraph starting with "In addition..."
    work_paras = [
        p for p in work_paras
        if not p.lower().startswith("in addition")
    ]

    if len(work_paras) >= 2:
        return clean_text(work_paras[1])
    elif len(work_paras) == 1:
        return clean_text(work_paras[0])
    return ""


def extract_required_years(work_experience_description):
    """
    Extract years from:
      A minimum of seven years ...
      A minimum of 7 years ...
    If text says no experience is required, return 0.
    """
    text = clean_text(work_experience_description)

    if re.search(r"not required to have professional work experience|no professional work experience", text,
                 re.IGNORECASE):
        return "0"

    number_words = {
        "zero": "0",
        "one": "1",
        "two": "2",
        "three": "3",
        "four": "4",
        "five": "5",
        "six": "6",
        "seven": "7",
        "eight": "8",
        "nine": "9",
        "ten": "10",
        "eleven": "11",
        "twelve": "12",
        "thirteen": "13",
        "fourteen": "14",
        "fifteen": "15",
    }

    match = re.search(
        r"A minimum of\s+(\d+|zero|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen)\s+years?",
        text,
        re.IGNORECASE
    )

    if match:
        value = match.group(1).lower()
        return number_words.get(value, value)

    return ""


def extract_specialties(all_text):
    """
    Extract specialties from GJP document.

    Priority 1:
        Lines like:
        3. Specialty: Legal Adviser
        Specialty: Political Affairs
        -> Legal Adviser

    Priority 2:
        Text like:
        includes the specialties "A", "B", "C"
        -> A, B, C

    If no specialties are found, return empty list.
    """

    text = all_text.replace("\xa0", " ")
    lines = [clean_text(line) for line in text.split("\n") if clean_text(line)]

    specialties = []

    # Priority 1: extract from "Specialty: xxx"
    for line in lines:
        match = re.search(
            r"(?:^\d+\.\s*)?Specialty\s*:\s*(.+)$",
            line,
            re.IGNORECASE
        )
        if match:
            specialty = clean_text(match.group(1))
            specialty = specialty.replace('"', "").replace("“", "").replace("”", "").strip()

            # Avoid accidental long sentence capture
            specialty = re.split(r"\s{2,}|Introduction|Responsibilities", specialty, flags=re.IGNORECASE)[0]
            specialty = clean_text(specialty)

            if specialty and specialty not in specialties:
                specialties.append(specialty)

    if specialties:
        return specialties

    # Priority 2: extract quoted specialties from "includes the specialties ..."
    match = re.search(
        r"includes\s+the\s+specialties\s+(.+?)(?:\.|$)",
        clean_text(all_text),
        re.IGNORECASE
    )

    if match:
        segment = match.group(1)

        quoted = re.findall(r'"([^"]+)"', segment)
        if quoted:
            return [clean_text(x) for x in quoted if clean_text(x)]

        segment = segment.replace('"', "").replace("“", "").replace("”", "")
        items = re.split(r"\s*[,;]\s*", segment)
        return [clean_text(x) for x in items if clean_text(x)]

    return []


def normalize_ksa_text(text):
    """
    Clean KSA text.
    Remove weird line breaks, repeated spaces, and table artifacts.
    """
    text = clean_text(text)
    text = text.replace("\u200b", "")
    text = text.replace("\ufeff", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def classify_ksa_type(ksa_text):
    """
    Classify KSA Type based on the extracted KSA description.

    Rules:
    - contains knowledge -> Knowledge
    - contains skill or skills -> Skills
    - contains ability or abilities -> Abilities
    - otherwise -> Skills
    """
    text = ksa_text.lower()

    if re.search(r"\bknowledge\b", text):
        return "Knowledge"

    if re.search(r"\bskills?\b", text):
        return "Skills"

    if re.search(r"\babilities?\b|\bability\b", text):
        return "Abilities"

    return "Skills"


def is_ksa_table(table):
    """
    Identify KSA table by checking whether the first row / first cell
    contains 'Knowledge, Skills, and Abilities' or similar wording.
    """
    if not table.rows:
        return False

    first_row_text = " ".join(
        clean_text(cell.text) for cell in table.rows[0].cells
    ).lower()

    first_cell_text = clean_text(table.rows[0].cells[0].text).lower()

    keywords = [
        "knowledge, skills, and abilities",
        "knowledge, skills, abilities",
        "knowledge skills and abilities",
        "knowledge skills abilities",
    ]

    return any(k in first_row_text for k in keywords) or any(k in first_cell_text for k in keywords)


def extract_ksas_from_docx(file_name, file_bytes):
    """
    Extract KSA titles from all KSA tables in one Word document.

    It does NOT distinguish specialties.
    If one document has several specialties, it extracts all KSA tables.
    Only the first column of each KSA table is used.
    """
    doc = Document(io.BytesIO(file_bytes))
    ksa_items = []

    for table in doc.tables:
        if not is_ksa_table(table):
            continue

        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            if not cells:
                continue

            first_col_value = normalize_ksa_text(cells[0].text)

            # skip header row
            if row_idx == 0:
                continue

            if not first_col_value:
                continue

            header_like = first_col_value.lower()
            if "knowledge, skills" in header_like:
                continue

            # skip notes
            if header_like.startswith("*note") or header_like.startswith("note:"):
                continue

            ksa_items.append({
                "ksa_text": first_col_value,
                "source_file": file_name,
            })

    return ksa_items


def build_ksa_dataframe(all_ksa_items):
    """
    Build unique KSA upload table according to the KSA Upload template column order.

    Template column order:
    1. KSA ID
    2. Title/ Short Description
    3. Title/ Short Description (French)
    4. Status
    5. KSA Type
    6. Proficiency Leveling Enabled
    7. Category 1
    8. Category 2
    9. Description
    10. Description (French)

    Also returns a separate dataframe for KSA descriptions over 140 characters.
    """
    unique_map = {}

    for item in all_ksa_items:
        ksa_text = normalize_ksa_text(item["ksa_text"])

        if not ksa_text:
            continue

        # Unique by case-insensitive normalized text
        unique_key = re.sub(r"\s+", " ", ksa_text.lower()).strip()

        if unique_key not in unique_map:
            unique_map[unique_key] = ksa_text

    rows = []

    for ksa_text in unique_map.values():
        row = {
            "KSA ID": "",
            "Title/ Short Description": ksa_text,
            "Title/ Short Description (French)": ksa_text,
            "Status": "Active",
            "KSA Type": classify_ksa_type(ksa_text),
            "Proficiency Leveling Enabled": "No",
            "Category 1": "test",
            "Category 2": "test",
            "Description": ksa_text,
            "Description (French)": ksa_text,
        }
        rows.append(row)

    KSA_COLUMNS = [
        "KSA ID",
        "Title/ Short Description",
        "Title/ Short Description (French)",
        "Status",
        "KSA Type",
        "Proficiency Leveling Enabled",
        "Category 1",
        "Category 2",
        "Description",
        "Description (French)",
    ]

    ksa_df = pd.DataFrame(rows)

    for col in KSA_COLUMNS:
        if col not in ksa_df.columns:
            ksa_df[col] = ""

    ksa_df = ksa_df[KSA_COLUMNS]

    # Auto-generate KSA ID from 1 to last row
    ksa_df["KSA ID"] = range(1, len(ksa_df) + 1)

    # Separate table for KSA short descriptions over 140 characters
    if len(ksa_df) > 0:
        over_140_df = ksa_df[
            ksa_df["Title/ Short Description"].astype(str).str.len() > 140
        ].copy()
    else:
        over_140_df = pd.DataFrame(columns=KSA_COLUMNS)

    return ksa_df, over_140_df


def format_excel_header(ws):
    """
    Make header row bold and add black border.
    """
    from openpyxl.styles import Font, Border, Side, Alignment

    black_side = Side(style="thin", color="000000")

    header_border = Border(
        left=black_side,
        right=black_side,
        top=black_side,
        bottom=black_side
    )

    for cell in ws[1]:
        cell.font = Font(bold=True, color="000000")
        cell.border = header_border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def create_ksa_excel_output(ksa_df, over_140_df):
    """
    Create one Excel file with two sheets:
    1. KSA Upload Template
    2. KSA Over 140 Characters

    Header row will be bold with black borders.
    """
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        ksa_df.to_excel(writer, index=False, sheet_name="KSA Upload Template")
        over_140_df.to_excel(writer, index=False, sheet_name="KSA Over 140 Characters")

        # Apply formatting to both sheets
        workbook = writer.book

        for sheet_name in ["KSA Upload Template", "KSA Over 140 Characters"]:
            ws = workbook[sheet_name]
            format_excel_header(ws)

            # Optional: adjust column width
            for col in ws.columns:
                max_length = 0
                col_letter = col[0].column_letter

                for cell in col:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))

                ws.column_dimensions[col_letter].width = min(max_length + 2, 50)

    output.seek(0)
    return output


def format_header_row(ws):
    """
    Make header row bold with black border.
    """
    black_side = Side(style="thin", color="000000")
    border = Border(left=black_side, right=black_side, top=black_side, bottom=black_side)

    for cell in ws[1]:
        cell.font = Font(bold=True, color="000000")
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def format_responsibilities_sheet(ws, bold_over_140=False):
    """
    Format Responsibilities sheet.
    If bold_over_140=True, bold KSA cells over 140 chars.
    """
    format_header_row(ws)

    # Basic formatting
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    # Column widths
    widths = {
        1: 16,  # Responsibility ID
        2: 35,  # Theme/Area
        3: 18,  # Status
        4: 70,  # Description
        5: 70,  # Description French
        6: 70,  # Mapped KSA
    }

    for col_idx, width in widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    # Bold KSA cells over 140 chars in column F
    if bold_over_140:
        for row_idx in range(2, ws.max_row + 1):
            cell = ws.cell(row=row_idx, column=6)
            if cell.value and len(str(cell.value)) > 140:
                cell.font = Font(bold=True)


def create_responsibilities_excel_output(resp_df, resp_over_140_df):
    """
    Create Excel with two sheets:
    1. Responsibilities Upload Template
    2. Responsibilities Over 140 KSA
    """
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        resp_df.to_excel(writer, index=False, sheet_name="Responsibilities Upload")
        resp_over_140_df.to_excel(writer, index=False, sheet_name="KSA Over 140")

        workbook = writer.book

        ws1 = workbook["Responsibilities Upload"]
        ws2 = workbook["KSA Over 140"]

        format_responsibilities_sheet(ws1, bold_over_140=False)
        format_responsibilities_sheet(ws2, bold_over_140=True)

    output.seek(0)
    return output


RESP_COLUMNS = [
    "Responsibility ID\n{Auto-generated}",
    "Theme/Area\n{Development of Products/tools, Collaboration/Coordination, Research/Study, Training, Implementation}",
    "Status \n{Active, Inactive}",
    "Description",
    "Description (French)",
    "Mapped KSA Short Description",
]


def iter_block_items(parent):
    """
    Yield paragraphs and tables in document order.
    This is important because each specialty has:
    Specialty heading -> Responsibilities table -> KSA table
    """
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def get_table_matrix(table):
    """
    Convert Word table into a clean matrix.
    """
    matrix = []

    for row in table.rows:
        row_values = [clean_text(cell.text) for cell in row.cells]
        if any(row_values):
            matrix.append(row_values)

    return matrix


def is_responsibility_table(table):
    """
    Identify Responsibilities table.

    Robust logic:
    A responsibilities table usually has:
    - Area of Work
    - No. / No
    - A third column containing job title, officer title, or responsibilities

    Some P-1/P-2 tables do NOT include the word 'Responsibilities' in the header,
    so we cannot require 'responsibilit' in header_text.
    """
    matrix = get_table_matrix(table)
    if not matrix:
        return False

    header = [clean_text(x).lower() for x in matrix[0]]
    header_text = " ".join(header)

    has_area = "area of work" in header_text

    has_no = any(
        re.fullmatch(r"no\.?", re.sub(r"\s+", " ", h).strip()) or
        re.sub(r"\s+", " ", h).strip() == "no"
        for h in header
    )

    # Avoid confusing KSA tables with responsibility tables
    is_ksa_like = (
        "knowledge" in header_text
        and "skills" in header_text
        and "abilities" in header_text
    )

    if is_ksa_like:
        return False

    # Main rule: Area of Work + No. is enough
    # because P-1/P-2 headers may only say "P-1 Assistant Human Rights Officer"
    if has_area and has_no:
        return True

    return False


def is_ksa_table_for_mapping(table):
    """
    Identify KSA mapping table.

    Typical header:
    Knowledge, Skills, and Abilities | P-2 ... Responsibilities
    """
    matrix = get_table_matrix(table)
    if not matrix:
        return False

    header_text = " ".join(matrix[0]).lower()

    return (
        "knowledge" in header_text
        and "skills" in header_text
        and "abilities" in header_text
        and "responsibilit" in header_text
    )


def parse_responsibility_table(table):
    """
    Parse Responsibilities table into:
        {
            responsibility_number: {
                "area_of_work": area,
                "description": desc
            }
        }

    Robust logic:
    - Detect Area of Work column
    - Detect No. column
    - Remove Area of Work and No. cells
    - Use the remaining meaningful text as responsibility description
    - Handles merged cells and blank Area of Work by fill-down
    """
    matrix = get_table_matrix(table)
    responsibilities = {}

    if not matrix or len(matrix) < 2:
        return responsibilities

    header = [clean_text(x).lower() for x in matrix[0]]

    area_col = None
    no_col = None

    for i, h in enumerate(header):
        h_clean = re.sub(r"\s+", " ", h).strip()

        if "area of work" in h_clean:
            area_col = i

        if re.fullmatch(r"no\.?", h_clean) or h_clean == "no":
            no_col = i

    # fallback
    if area_col is None:
        area_col = 0
    if no_col is None:
        no_col = 1

    last_area_of_work = ""

    for row in matrix[1:]:
        if not row:
            continue

        row = [clean_text(x) for x in row]

        # make row long enough
        max_col = max(area_col, no_col)
        while len(row) <= max_col:
            row.append("")

        area = clean_text(row[area_col])
        no_text = clean_text(row[no_col])

        # Fill down Area of Work when merged/blank
        if area:
            last_area_of_work = area
        else:
            area = last_area_of_work

        # No. must be a number
        no_match = re.search(r"\b\d+\b", no_text)

        # if No. column fails, scan cells for a pure number
        if not no_match:
            for cell in row:
                cell_text = clean_text(cell)
                if re.fullmatch(r"\d+", cell_text):
                    no_match = re.search(r"\d+", cell_text)
                    break

        if not no_match:
            continue

        no = int(no_match.group(0))

        # Remove Area of Work and No. columns, remaining cells are potential descriptions
        desc_candidates = []
        for i, cell in enumerate(row):
            if i in [area_col, no_col]:
                continue

            cell = clean_text(cell)

            if not cell:
                continue

            lower = cell.lower()

            # skip header-like cells
            if "area of work" in lower:
                continue
            if re.fullmatch(r"no\.?", lower):
                continue
            if lower.startswith("*note") or lower.startswith("note:"):
                continue

            # skip pure number cells
            if re.fullmatch(r"\d+", cell):
                continue

            desc_candidates.append(cell)

        if not desc_candidates:
            continue

        # Responsibility description is usually the longest remaining cell
        desc = max(desc_candidates, key=len)
        desc = clean_text(desc)

        if not desc:
            continue

        responsibilities[no] = {
            "area_of_work": area,
            "description": desc,
        }

    return responsibilities


def extract_numbers_from_text(text):
    """
    Extract responsibility numbers from KSA mapping cells.
    """
    text = clean_text(text)
    nums = re.findall(r"\b\d+\b", text)
    return [int(x) for x in nums]


def parse_ksa_mapping_table(table):
    """
    Parse KSA table into mapping:
        responsibility_number -> [ksa1, ksa2, ...]

    KSA table:
        first column = KSA short description
        other columns/cells = responsibility numbers
    """
    matrix = get_table_matrix(table)
    mapping = {}

    if not matrix:
        return mapping

    for row_idx, row in enumerate(matrix[1:], start=1):
        if not row:
            continue

        ksa_text = normalize_ksa_text(row[0])

        if not ksa_text:
            continue

        lower = ksa_text.lower()

        # Skip note rows and header-like rows
        if lower.startswith("*note") or lower.startswith("note:"):
            continue

        if "knowledge, skills" in lower:
            continue

        number_text = " ".join(row[1:])
        resp_numbers = extract_numbers_from_text(number_text)

        for no in resp_numbers:
            mapping.setdefault(no, [])
            if ksa_text not in mapping[no]:
                mapping[no].append(ksa_text)

    return mapping


def extract_specialty_from_paragraph(text):
    """
    Extract specialty from paragraph like:
        1. Specialty: General
        3. Specialty: Legal Adviser
    """
    text = clean_text(text)

    match = re.search(
        r"(?:^\d+\.\s*)?Specialty\s*:\s*(.+)$",
        text,
        re.IGNORECASE
    )

    if not match:
        return None

    specialty = clean_text(match.group(1))
    specialty = specialty.replace('"', "").replace("“", "").replace("”", "").strip()
    return specialty


def build_responsibility_rows_for_one_section(
    responsibilities,
    ksa_mapping,
    specialty_name,
    folder_name,
    paragraphs,
    seen_resp_keys=None
):
    """
    Build upload rows for one specialty section.

    Theme/Area should be Area of Work from the Responsibilities table,
    NOT specialty.

    Deduplication rule:
    Area of Work + Responsibility Description must be unique.
    Punctuation-only differences are treated as the same.

    First KSA row contains full responsibility data.
    Additional mapped KSAs only fill Mapped KSA Short Description.
    """
    rows = []

    if seen_resp_keys is None:
        seen_resp_keys = set()

    for no in sorted(responsibilities.keys()):
        area_of_work = clean_text(responsibilities[no].get("area_of_work", ""))
        resp_desc = clean_text(responsibilities[no].get("description", ""))

        unique_key = (
            normalize_for_unique(area_of_work),
            normalize_for_unique(resp_desc)
        )

        if unique_key in seen_resp_keys:
            continue

        seen_resp_keys.add(unique_key)

        mapped_ksas = ksa_mapping.get(no, [])

        if not mapped_ksas:
            mapped_ksas = [""]

        has_over_140 = any(len(str(ksa)) > 140 for ksa in mapped_ksas)

        for idx, ksa in enumerate(mapped_ksas):
            if idx == 0:
                row = {
                    "Responsibility ID\n{Auto-generated}": "",
                    "Theme/Area\n{Development of Products/tools, Collaboration/Coordination, Research/Study, Training, Implementation}": area_of_work,
                    "Status \n{Active, Inactive}": "Active",
                    "Description": resp_desc,
                    "Description (French)": resp_desc,
                    "Mapped KSA Short Description": ksa,
                    "_resp_no": no,
                    "_over_140": has_over_140,
                    "_unique_key": unique_key,
                }
            else:
                row = {
                    "Responsibility ID\n{Auto-generated}": "",
                    "Theme/Area\n{Development of Products/tools, Collaboration/Coordination, Research/Study, Training, Implementation}": "",
                    "Status \n{Active, Inactive}": "",
                    "Description": "",
                    "Description (French)": "",
                    "Mapped KSA Short Description": ksa,
                    "_resp_no": no,
                    "_over_140": has_over_140,
                    "_unique_key": unique_key,
                }

            rows.append(row)

    return rows


def extract_responsibilities_with_ksa_mapping_from_docx(file_name, file_bytes, folder_name):
    """
    Extract Responsibilities upload rows from one DOCX.

    New robust logic:
    - Read every Responsibilities table.
    - Find the nearest following KSA table for mapping.
    - If there is no specialty heading, still process responsibilities.
    - If no KSA table is found, still output responsibilities with blank mapped KSA.
    - Theme/Area comes from Area of Work column.
    - Deduplicate by Area of Work + Responsibility Description.
    """
    doc = Document(io.BytesIO(file_bytes))

    paragraphs_for_fallback = []
    current_specialty = ""

    all_rows = []
    seen_resp_keys = set()

    pending_responsibilities = None
    pending_specialty = ""

    def flush_pending(ksa_mapping=None):
        """
        Output pending responsibilities.
        If ksa_mapping is None or empty, Mapped KSA Short Description will be blank.
        """
        nonlocal pending_responsibilities, pending_specialty, all_rows, seen_resp_keys

        if not pending_responsibilities:
            return

        if ksa_mapping is None:
            ksa_mapping = {}

        section_rows = build_responsibility_rows_for_one_section(
            responsibilities=pending_responsibilities,
            ksa_mapping=ksa_mapping,
            specialty_name=pending_specialty,
            folder_name=folder_name,
            paragraphs=paragraphs_for_fallback,
            seen_resp_keys=seen_resp_keys
        )

        all_rows.extend(section_rows)

        pending_responsibilities = None
        pending_specialty = ""

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = clean_text(block.text)

            if text:
                paragraphs_for_fallback.append(text)

            specialty = extract_specialty_from_paragraph(text)
            if specialty:
                current_specialty = specialty

        elif isinstance(block, Table):

            # 1. If this is a Responsibilities table, store it.
            # If another responsibility table is already pending, flush it first.
            if is_responsibility_table(block):
                print("Found responsibility table in:", file_name)

                flush_pending(ksa_mapping={})

                parsed_resp = parse_responsibility_table(block)

                if parsed_resp:
                    pending_responsibilities = parsed_resp
                    pending_specialty = current_specialty

            # 2. If this is a KSA table and we have pending responsibilities,
            # use this nearest KSA table to map KSAs.
            elif is_ksa_table_for_mapping(block):
                if pending_responsibilities:
                    ksa_mapping = parse_ksa_mapping_table(block)
                    flush_pending(ksa_mapping=ksa_mapping)

                # If no pending responsibilities, ignore this KSA table.

    # 3. End of document: still output any pending responsibilities.
    flush_pending(ksa_mapping={})

    return all_rows


def build_responsibilities_dataframe(all_resp_rows):
    """
    Build final Responsibilities dataframe and over-140 dataframe.
    """
    df = pd.DataFrame(all_resp_rows)

    if df.empty:
        df = pd.DataFrame(columns=RESP_COLUMNS + ["_resp_no", "_over_140"])

    for col in RESP_COLUMNS:
        if col not in df.columns:
            df[col] = ""

    # Auto-generate Responsibility ID only on rows with actual responsibility description
    # KSA-only continuation rows stay blank.
    resp_id = 1
    for idx in df.index:
        if clean_text(df.at[idx, "Description"]):
            df.at[idx, "Responsibility ID\n{Auto-generated}"] = str(resp_id)
            resp_id += 1
        else:
            df.at[idx, "Responsibility ID\n{Auto-generated}"] = ""

    # Sheet 2: copy entire responsibility group if any mapped KSA is over 140 chars
    if "_over_140" in df.columns:
        over_140_df = df[df["_over_140"] == True].copy()
    else:
        over_140_df = pd.DataFrame(columns=df.columns)

    # Keep final visible columns
    visible_df = df[RESP_COLUMNS].copy()
    visible_over_140_df = over_140_df[RESP_COLUMNS].copy()

    return visible_df, visible_over_140_df


SPECIALTY_COLUMNS = [
    "Specialty ID",
    "Status",
    "Specialty Name",
    "Specialty Name (French)",
    "Description",
    "Description (French)",
    "Certficate Name",
    "Mapped Responsibility Description",
    "Responsibility Mandatory\n{Yes, No}",
]


def paragraph_has_any_bold(paragraph):
    """
    Return True if any run in the paragraph is bold.
    """
    for run in paragraph.runs:
        if run.bold:
            return True
    return False


def cell_has_bold_text(cell):
    """
    Return True if a Word table cell contains bold text.
    Used to determine Responsibility Mandatory {Yes, No}.
    """
    for paragraph in cell.paragraphs:
        if paragraph_has_any_bold(paragraph):
            return True
    return False


def extract_specialty_name_from_intro_title(paragraphs):
    """
    For GJPs without explicit 'Specialty:' heading, e.g. P-1:
        The Assistant Human Rights Officer P-1 GJP is available for ...
    Extract:
        Assistant Human Rights Officer

    It extracts text after 'The' and before level.
    """
    level_pattern = r"(?:P-[1-5]|D-1|NO-[A-D]|NO-[A-D])"

    for txt in paragraphs:
        text = clean_text(txt)

        match = re.search(
            rf"^The\s+(.+?)\s+{level_pattern}\s+GJP\s+is\s+available\s+for\b",
            text,
            re.IGNORECASE
        )

        if match:
            return clean_text(match.group(1))

    return ""


def extract_intro_text_until_responsibilities(blocks, start_index):
    """
    Starting from a Specialty heading or fallback start, find the Introduction section
    and collect all paragraph text until Responsibilities.
    """
    intro_started = False
    intro_parts = []

    for block in blocks[start_index + 1:]:
        if isinstance(block, Paragraph):
            text = clean_text(block.text)
            if not text:
                continue

            lower = text.lower()

            if lower == "introduction":
                intro_started = True
                continue

            if intro_started and lower == "responsibilities":
                break

            if intro_started:
                # Stop if next specialty appears unexpectedly
                if extract_specialty_from_paragraph(text):
                    break
                intro_parts.append(text)

        elif isinstance(block, Table):
            # If we hit a table after intro started, intro is done
            if intro_started:
                break

    return clean_text(" ".join(intro_parts))


def find_fallback_intro_start_index(blocks):
    """
    For documents without explicit Specialty heading, such as P-1,
    use the first Introduction section as the specialty description source.
    """
    for i, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            if clean_text(block.text).lower() == "introduction":
                return i - 1 if i > 0 else 0
    return 0


def extract_responsibility_rows_with_bold(table):
    """
    Extract responsibilities from a responsibility table.

    Returns a list of:
        {
            "description": responsibility description,
            "mandatory": "Yes" / "No"
        }

    Logic:
    - Find Area of Work and No. columns.
    - Remove those two columns.
    - Use remaining longest cell as responsibility description.
    - Check whether the selected description cell has bold text.
    """
    matrix = get_table_matrix(table)

    if not matrix or len(matrix) < 2:
        return []

    header = [clean_text(x).lower() for x in matrix[0]]

    area_col = None
    no_col = None

    for i, h in enumerate(header):
        h_clean = re.sub(r"\s+", " ", h).strip()

        if "area of work" in h_clean:
            area_col = i

        if re.fullmatch(r"no\.?", h_clean) or h_clean == "no":
            no_col = i

    if area_col is None:
        area_col = 0
    if no_col is None:
        no_col = 1

    results = []

    for row in table.rows[1:]:
        cells = row.cells
        cell_texts = [clean_text(cell.text) for cell in cells]

        if not any(cell_texts):
            continue

        max_col = max(area_col, no_col)
        while len(cell_texts) <= max_col:
            cell_texts.append("")

        no_text = clean_text(cell_texts[no_col])
        no_match = re.search(r"\b\d+\b", no_text)

        if not no_match:
            # scan for a pure number if No column failed
            for cell_text in cell_texts:
                if re.fullmatch(r"\d+", clean_text(cell_text)):
                    no_match = re.search(r"\d+", clean_text(cell_text))
                    break

        if not no_match:
            continue

        desc_candidates = []

        for i, cell_text in enumerate(cell_texts):
            if i in [area_col, no_col]:
                continue

            text = clean_text(cell_text)
            if not text:
                continue

            lower = text.lower()

            if lower.startswith("*note") or lower.startswith("note:"):
                continue
            if "area of work" in lower:
                continue
            if re.fullmatch(r"no\.?", lower):
                continue
            if re.fullmatch(r"\d+", text):
                continue

            desc_candidates.append((i, text))

        if not desc_candidates:
            continue

        # Choose the longest remaining cell as responsibility description
        desc_col, desc_text = max(desc_candidates, key=lambda x: len(x[1]))

        mandatory = "No"
        if desc_col < len(cells):
            if cell_has_bold_text(cells[desc_col]):
                mandatory = "Yes"

        results.append({
            "description": desc_text,
            "mandatory": mandatory,
        })

    return results


def extract_specialties_upload_rows_from_docx(file_name, file_bytes, folder_name):
    """
    Extract Specialties Upload Template rows from one DOCX.

    Required:
    - Status = Active
    - Specialty Name:
        If document has 'Specialty: X', use X, with General -> Generalist - Human rights
        If no 'Specialty:', use text after 'The' and before level in:
            The Assistant Human Rights Officer P-1 GJP is available for ...
    - Specialty Name (French) = Specialty Name
    - Description = Introduction content under corresponding specialty
    - Description (French) = Description
    - Certficate Name = blank
    - Mapped Responsibility Description = each responsibility description
    - Responsibility Mandatory {Yes, No} = Yes if responsibility description cell has bold text, else No
    """
    doc = Document(io.BytesIO(file_bytes))
    blocks = list(iter_block_items(doc))

    # For fallback title extraction, collect normal paragraphs
    paragraph_texts = [
        clean_text(block.text)
        for block in blocks
        if isinstance(block, Paragraph) and clean_text(block.text)
    ]

    rows = []

    current_specialty_raw = None
    current_specialty_name = None
    current_intro = ""

    found_any_specialty_heading = False

    for idx, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            text = clean_text(block.text)
            specialty_raw = extract_specialty_from_paragraph(text)

            if specialty_raw:
                found_any_specialty_heading = True
                current_specialty_raw = specialty_raw
                current_specialty_name = clean_mapped_specialty_name(
                    specialty=specialty_raw,
                    folder_name=folder_name,
                    paragraphs=paragraph_texts
                )
                current_intro = extract_intro_text_until_responsibilities(blocks, idx)

        elif isinstance(block, Table):
            if is_responsibility_table(block):
                responsibility_rows = extract_responsibility_rows_with_bold(block)

                if not responsibility_rows:
                    continue

                # If no explicit Specialty heading exists, use fallback specialty name and introduction
                if not current_specialty_name:
                    fallback_name = extract_specialty_name_from_intro_title(paragraph_texts)
                    if not fallback_name:
                        fallback_name = extract_gjp_job_title(paragraph_texts)

                    current_specialty_name = fallback_name
                    fallback_start = find_fallback_intro_start_index(blocks)
                    current_intro = extract_intro_text_until_responsibilities(blocks, fallback_start)

                # Build rows: first responsibility row has specialty info;
                # following mapped responsibilities only fill mapped responsibility + mandatory.
                for r_idx, resp in enumerate(responsibility_rows):
                    if r_idx == 0:
                        row = {
                            "Specialty ID": "",
                            "Status": "Active",
                            "Specialty Name": current_specialty_name,
                            "Specialty Name (French)": current_specialty_name,
                            "Description": current_intro,
                            "Description (French)": current_intro,
                            "Certficate Name": "",
                            "Mapped Responsibility Description": resp["description"],
                            "Responsibility Mandatory\n{Yes, No}": resp["mandatory"],
                        }
                    else:
                        row = {
                            "Specialty ID": "",
                            "Status": "",
                            "Specialty Name": "",
                            "Specialty Name (French)": "",
                            "Description": "",
                            "Description (French)": "",
                            "Certficate Name": "",
                            "Mapped Responsibility Description": resp["description"],
                            "Responsibility Mandatory\n{Yes, No}": resp["mandatory"],
                        }

                    rows.append(row)

    return rows

def build_specialties_dataframe(all_specialty_rows):
    """
    Build final Specialties dataframe with exact template column order.
    Specialty ID is auto-generated only on rows with actual Specialty Name.
    """
    df = pd.DataFrame(all_specialty_rows)

    if df.empty:
        df = pd.DataFrame(columns=SPECIALTY_COLUMNS)

    for col in SPECIALTY_COLUMNS:
        if col not in df.columns:
            df[col] = ""

    df = df[SPECIALTY_COLUMNS].copy()

    # Auto-generate Specialty ID only on first row of each specialty
    specialty_id = 1
    for idx in df.index:
        if clean_text(df.at[idx, "Specialty Name"]):
            df.at[idx, "Specialty ID"] = str(specialty_id)
            specialty_id += 1
        else:
            df.at[idx, "Specialty ID"] = ""

    return df


def format_specialties_sheet(ws):
    """
    Format Specialties Excel output.
    """
    format_header_row(ws)

    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = {
        1: 14,  # Specialty ID
        2: 14,  # Status
        3: 28,  # Specialty Name
        4: 28,  # Specialty Name French
        5: 70,  # Description
        6: 70,  # Description French
        7: 22,  # Certficate Name
        8: 75,  # Mapped Responsibility Description
        9: 22,  # Responsibility Mandatory
    }

    for col_idx, width in widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = width


def create_specialties_excel_output(specialties_df):
    """
    Create Specialties Excel output.
    """
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        specialties_df.to_excel(writer, index=False, sheet_name="Specialties Upload")

        workbook = writer.book
        ws = workbook["Specialties Upload"]
        format_specialties_sheet(ws)

    output.seek(0)
    return output


def extract_one_docx(file_name, file_bytes, folder_name):
    """
    Extract one GJP document into one or multiple rows.
    If specialties exist, put each specialty in a separate row.

    Column names strictly follow the GJP Upload template.
    French fields are temporarily filled with the same English content.
    """
    paragraphs, all_text, tables_text = get_docx_text_and_tables(file_bytes)

    gjp_title = extract_gjp_job_title(paragraphs)
    job_code = extract_job_code(all_text)
    description = extract_description(paragraphs)

    edu_desc = extract_education_description(paragraphs)
    min_edu = extract_minimum_education_level(edu_desc)

    work_desc = extract_work_experience_description(paragraphs)
    required_years = extract_required_years(work_desc)

    specialties = extract_specialties(all_text)

    # If no specialty found, still generate one row with blank Mapped Specialty Name
    if not specialties:
        specialties = [""]

    rows = []

    for specialty in specialties:
        row = {
            "GJP Template ID": "",
            "Status": "Active",
            "GJP Area/ Workstream": folder_name,

            "GJP Job Title": gjp_title,
            "GJP Job Title (French)": gjp_title,

            "Description": description,
            "Description (French)": description,

            "Jobcode ": job_code,

            "Required Years of Work Experience": required_years,

            "Work Experience Description ": work_desc,
            "Work Experience Description (French)": work_desc,

            "Minimum Level of Education Required": min_edu,

            "Education Description (English)": edu_desc,
            "Education Description (French)": edu_desc,

            "Mapped Specialty Name": clean_mapped_specialty_name(specialty=specialty, folder_name=folder_name,
                                                                 paragraphs=paragraphs),
        }

        rows.append(row)

    return rows


def read_uploaded_files(uploaded_files):
    """
    Accept multiple uploaded files.
    Supports:
      - .docx files selected directly
      - .zip file containing docx files
    """
    docx_files = []

    for uploaded in uploaded_files:
        file_name = uploaded.name
        file_bytes = uploaded.read()

        if file_name.lower().endswith(".docx"):
            docx_files.append((file_name, file_bytes))

        elif file_name.lower().endswith(".zip"):
            with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as z:
                for name in z.namelist():
                    if name.lower().endswith(".docx") and not name.startswith("__MACOSX"):
                        docx_files.append((Path(name).name, z.read(name)))

    return docx_files


def fill_template_if_provided(df, template_file):
    """
    If user uploads an Excel template, fill the columns that match.
    Otherwise, return a normal generated Excel.
    """
    output = io.BytesIO()

    if template_file is None:
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="GJP Upload")
        output.seek(0)
        return output

    template_bytes = template_file.read()
    wb = load_workbook(io.BytesIO(template_bytes))
    ws = wb.active

    # Find header row by searching for known columns
    known_cols = set(df.columns)
    header_row = None
    header_map = {}

    for row in range(1, min(ws.max_row, 20) + 1):
        temp_map = {}
        for col in range(1, ws.max_column + 1):
            value = ws.cell(row=row, column=col).value
            if value:
                header = clean_text(value)
                temp_map[header] = col

        matched = len(set(temp_map.keys()) & known_cols)
        if matched >= 3:
            header_row = row
            header_map = temp_map
            break

    if header_row is None:
        # fallback: create a new sheet
        ws = wb.create_sheet("Generated GJP Upload")
        for col_idx, col_name in enumerate(df.columns, start=1):
            ws.cell(row=1, column=col_idx).value = col_name
        header_row = 1
        header_map = {col: idx for idx, col in enumerate(df.columns, start=1)}

    start_row = header_row + 1

    # clear existing values under matching columns
    for excel_header, col_idx in header_map.items():
        if excel_header in df.columns:
            for row in range(start_row, ws.max_row + 1):
                ws.cell(row=row, column=col_idx).value = None

    # write data
    for r_idx, record in enumerate(df.to_dict(orient="records"), start=start_row):
        for excel_header, col_idx in header_map.items():
            if excel_header in record:
                ws.cell(row=r_idx, column=col_idx).value = record[excel_header]

    wb.save(output)
    output.seek(0)
    return output


# =========================
# AoW/KSA mapping matrix generator
# =========================

AOW_RESPONSIBILITY_HEADERS = ("area of work", "no", "responsibilities")
AOW_KSA_HEADER_KEYWORDS = ("knowledge", "skills", "abilities")


def aow_row_texts(row):
    return [clean_text(cell.text) for cell in row.cells]


def aow_extract_numbers(text):
    return set(re.findall(r"\d+", clean_text(text)))


def aow_looks_like_responsibility_table(table):
    if not table.rows:
        return False
    header = " | ".join(aow_row_texts(table.rows[0])).lower()
    return all(keyword in header for keyword in AOW_RESPONSIBILITY_HEADERS)


def aow_looks_like_ksa_table(table):
    if not table.rows:
        return False
    first_cell = clean_text(table.rows[0].cells[0].text).lower()
    return all(keyword in first_cell for keyword in AOW_KSA_HEADER_KEYWORDS)


def aow_find_responsibility_table(document):
    candidates = [table for table in document.tables if aow_looks_like_responsibility_table(table)]
    if candidates:
        return candidates[0]
    raise ValueError(
        "No responsibilities table was found. The Word file should include "
        "Area of Work, No., and Responsibilities headers."
    )


def aow_find_ksa_table(document):
    candidates = [table for table in document.tables if aow_looks_like_ksa_table(table)]
    if candidates:
        return candidates[0]
    raise ValueError(
        "No KSA mapping table was found. The Word file should include a "
        "Knowledge, Skills, and Abilities table."
    )


def aow_header_index(headers, keywords):
    lowered = [header.lower() for header in headers]
    for keyword in keywords:
        for index, header in enumerate(lowered):
            if keyword in header:
                return index
    return None


def aow_parse_responsibilities(document):
    table = aow_find_responsibility_table(document)
    headers = aow_row_texts(table.rows[0])
    theme_col = aow_header_index(headers, ("area of work", "theme"))
    number_col = aow_header_index(headers, ("no.", "no", "nr", "number"))
    resp_col = aow_header_index(headers, ("responsibilities", "responsibility"))

    if theme_col is None or number_col is None or resp_col is None:
        raise ValueError(
            "The responsibilities table headers could not be recognized. "
            "Required headers: Area of Work, No., Responsibilities."
        )

    responsibilities = []
    last_theme = ""
    for row in table.rows[1:]:
        cells = aow_row_texts(row)
        if len(cells) <= max(theme_col, number_col, resp_col):
            continue

        theme = cells[theme_col] or last_theme
        number_match = re.search(r"\d+", cells[number_col])
        number = number_match.group(0) if number_match else clean_text(cells[number_col])
        text = cells[resp_col]

        if theme:
            last_theme = theme
        if number and text:
            responsibilities.append({
                "theme": theme,
                "number": number,
                "text": text,
            })

    if not responsibilities:
        raise ValueError("No valid responsibility rows were found.")
    return responsibilities


def aow_parse_ksa_mapping(document):
    table = aow_find_ksa_table(document)
    ksas = []
    mapping_by_number = {}

    for row in table.rows[1:]:
        cells = aow_row_texts(row)
        if not cells:
            continue

        ksa = cells[0]
        if not ksa or all(keyword in ksa.lower() for keyword in AOW_KSA_HEADER_KEYWORDS):
            continue

        if ksa not in ksas:
            ksas.append(ksa)

        for number in aow_extract_numbers(" ".join(cells[1:])):
            mapping_by_number.setdefault(number, set()).add(ksa)

    if not ksas:
        raise ValueError("No valid KSA rows were found.")
    return ksas, mapping_by_number


def aow_build_dataframe(responsibilities, ksas, mapping_by_number, mark, include_assigned_sme, collapse_theme):
    rows = []
    previous_theme = None

    for responsibility in responsibilities:
        shown_theme = responsibility["theme"]
        if collapse_theme and responsibility["theme"] == previous_theme:
            shown_theme = ""
        previous_theme = responsibility["theme"]

        row = {
            "Theme": shown_theme,
            "Number": responsibility["number"],
            "Responsibilities": responsibility["text"],
        }
        if include_assigned_sme:
            row["Assigned SME"] = ""

        linked_ksas = mapping_by_number.get(responsibility["number"], set())
        for ksa in ksas:
            row[ksa] = mark if ksa in linked_ksas else ""

        rows.append(row)

    base_columns = ["Theme", "Number", "Responsibilities"]
    if include_assigned_sme:
        base_columns.append("Assigned SME")

    return pd.DataFrame(rows, columns=base_columns + ksas)


def aow_apply_sheet_style(ws, row_count, col_count, include_assigned_sme):
    dark_fill = PatternFill("solid", fgColor="1F4E78")
    note_fill = PatternFill("solid", fgColor="EAF3F8")
    header_font = Font(color="FFFFFF", bold=True)
    note_font = Font(color="1F4E78", bold=True)
    thin_gray = Side(style="thin", color="D9E2EA")
    medium_blue = Side(style="medium", color="1F4E78")
    border = Border(left=thin_gray, right=thin_gray, top=thin_gray, bottom=thin_gray)

    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "E4" if include_assigned_sme else "D4"

    ws.cell(1, 1).value = "Please enter an X to map KSAs to responsibilities"
    ws.cell(2, 1).value = "Generated from Word responsibilities and KSA mapping tables"

    for row in (1, 2):
        for col in range(1, col_count + 1):
            cell = ws.cell(row, col)
            cell.fill = note_fill
            cell.font = note_font
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 22
    ws.row_dimensions[2].height = 22

    for cell in ws[3]:
        cell.fill = dark_fill
        cell.font = header_font
        cell.border = Border(top=medium_blue, bottom=medium_blue)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[3].height = 58

    theme_fills = [
        PatternFill("solid", fgColor="FFFFFF"),
        PatternFill("solid", fgColor="F6FAFD"),
    ]
    current_fill_index = 0
    last_visible_theme = None
    first_ksa_col = 5 if include_assigned_sme else 4

    for row in range(4, row_count + 4):
        visible_theme = clean_text(ws.cell(row, 1).value)
        if visible_theme and visible_theme != last_visible_theme:
            current_fill_index = 1 - current_fill_index
            last_visible_theme = visible_theme

        fill = theme_fills[current_fill_index]
        for col in range(1, col_count + 1):
            cell = ws.cell(row, col)
            cell.fill = fill
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

        ws.cell(row, 2).alignment = Alignment(horizontal="center", vertical="top", wrap_text=True)
        for col in range(first_ksa_col, col_count + 1):
            ws.cell(row, col).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.row_dimensions[row].height = 54

    widths = {1: 22, 2: 8, 3: 72}
    if include_assigned_sme:
        widths[4] = 16
    for col in range(first_ksa_col, col_count + 1):
        widths[col] = 24
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = width

    ws.auto_filter.ref = f"A3:{get_column_letter(col_count)}{row_count + 3}"


def aow_dataframe_to_xlsx(df, include_assigned_sme):
    workbook = Workbook()
    ws = workbook.active
    ws.title = "AoW Mapping"

    for col_index, column_name in enumerate(df.columns, start=1):
        ws.cell(3, col_index).value = column_name
    for row_index, values in enumerate(df.itertuples(index=False, name=None), start=4):
        for col_index, value in enumerate(values, start=1):
            ws.cell(row_index, col_index).value = value

    aow_apply_sheet_style(ws, len(df), len(df.columns), include_assigned_sme)

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def aow_convert_docx(file_bytes, mark, include_assigned_sme, collapse_theme):
    document = Document(io.BytesIO(file_bytes))
    responsibilities = aow_parse_responsibilities(document)
    ksas, mapping_by_number = aow_parse_ksa_mapping(document)
    df = aow_build_dataframe(
        responsibilities=responsibilities,
        ksas=ksas,
        mapping_by_number=mapping_by_number,
        mark=mark,
        include_assigned_sme=include_assigned_sme,
        collapse_theme=collapse_theme,
    )
    excel_output = aow_dataframe_to_xlsx(df, include_assigned_sme=include_assigned_sme)

    responsibility_numbers = {item["number"] for item in responsibilities}
    mapped_numbers = set(mapping_by_number)
    warnings = []

    unknown_numbers = sorted(mapped_numbers - responsibility_numbers, key=lambda x: int(x) if x.isdigit() else 99999)
    if unknown_numbers:
        warnings.append(f"The KSA table references responsibility numbers that were not found: {', '.join(unknown_numbers)}")

    unmapped_numbers = sorted(responsibility_numbers - mapped_numbers, key=lambda x: int(x) if x.isdigit() else 99999)
    if unmapped_numbers:
        warnings.append(f"These responsibility numbers do not have any KSA marks: {', '.join(unmapped_numbers)}")

    return df, excel_output, responsibilities, ksas, warnings


# =========================
# Streamlit website
# =========================

from gjp_document_generator import render_gjp_document_generator
from ksa_mapping_streamlit_app import render_ksa_conversion_tables


def render_aow_ksa_mapping_matrix():
    st.title("AoW/KSA Mapping Matrix Generator")
    st.write(
        "Upload one Word document with responsibilities and KSA mapping tables. "
        "This tool generates an Excel matrix with Theme, Number, Responsibilities, "
        "and one column for each KSA. Mapped cells are marked with a capital X."
    )

    with st.sidebar:
        st.markdown("---")
        st.header("AoW/KSA settings")
        include_assigned_sme = st.checkbox("Include blank Assigned SME column", value=True)
        collapse_theme = st.checkbox("Show each theme only once", value=True)
        st.caption("KSA mark: X")

    uploaded_file = st.file_uploader(
        "Upload a Word document",
        type=["docx"],
        accept_multiple_files=False
    )

    if uploaded_file is None:
        st.info("Upload a Word document to begin.")
        return

    try:
        df, excel_output, responsibilities, ksas, warnings = aow_convert_docx(
            file_bytes=uploaded_file.getvalue(),
            mark="X",
            include_assigned_sme=include_assigned_sme,
            collapse_theme=collapse_theme,
        )
    except Exception as e:
        st.error(f"Conversion failed: {e}")
        return

    col1, col2, col3 = st.columns(3)
    col1.metric("Responsibilities", len(responsibilities))
    col2.metric("KSA columns", len(ksas))
    col3.metric("Output columns", len(df.columns))

    for warning in warnings:
        st.warning(warning)

    st.subheader("Preview")
    st.dataframe(df, use_container_width=True, height=520)

    output_name = uploaded_file.name.rsplit(".", 1)[0] + "_AoW_KSA_mapping.xlsx"
    st.download_button(
        label="Download AoW/KSA Excel",
        data=excel_output,
        file_name=output_name,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


def render_word_to_excel():
    st.title("GJP Upload Template Generator")
    st.write(
        "Upload one department folder as a ZIP file, or select multiple DOCX files. "
        "The app will extract GJP fields and generate an Excel upload template."
    )

    folder_name = st.text_input(
        "GJP Area / workstream",
        value="Human Rights",
        help="This will be filled into the GJP Area/workstream column. Usually this is the department or folder name."
    )

    uploaded_files = st.file_uploader(
        "Upload DOCX files or one ZIP folder",
        type=["docx", "zip"],
        accept_multiple_files=True
    )

    st.markdown("---")
    st.header("Generate Job Information Upload Table")

    st.write(
        "This function extracts job information from all uploaded GJP Word documents and creates the job "
        "information upload workbook. The output includes GJP Template ID, Status, GJP Area/Workstream, "
        "English and French job titles and descriptions, Job Code, required years and descriptions of work "
        "experience, minimum education level, English and French education descriptions, and the mapped "
        "specialty name. French-language columns currently copy the English content and should be reviewed "
        "and translated before upload."
    )

    if st.button("Generate Excel"):
        if not uploaded_files:
            st.error("Please upload at least one DOCX file or one ZIP file.")
        else:
            all_rows = []
            docx_files = read_uploaded_files(uploaded_files)

            if not docx_files:
                st.error("No DOCX files found. Please upload DOCX files or a ZIP containing DOCX files.")
            else:
                for file_name, file_bytes in docx_files:
                    try:
                        rows = extract_one_docx(file_name, file_bytes, folder_name)
                        all_rows.extend(rows)
                    except Exception as e:
                        st.warning(f"Failed to process {file_name}: {e}")

                if not all_rows:
                    st.error("No data extracted.")
                else:
                    df = pd.DataFrame(all_rows)
                    df["GJP Template ID"] = range(1, len(df) + 1)

                    st.success(f"Extracted {len(df)} row(s) from {len(docx_files)} document(s).")
                    st.dataframe(df, use_container_width=True)

                    excel_output = fill_template_if_provided(df, None)

                    st.download_button(
                        label="Download generated Excel",
                        data=excel_output,
                        file_name="GJP_upload_generated.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

    st.markdown("---")
    st.header("Generate KSA Upload Table")

    st.write(
        "This function extracts unique KSAs from all uploaded GJP Word documents. "
        "It reads the first column of each Knowledge, Skills, Abilities (KSAs) table, "
        "fills the KSA upload table, and creates a separate sheet for KSAs over 140 characters."
    )

    if st.button("Generate KSA Table"):
        if not uploaded_files:
            st.error("Please upload at least one DOCX file or one ZIP file first.")
        else:
            all_ksa_items = []

            # Important: uploaded_files may already have been read once.
            # Reset file pointer before reading again.
            for uploaded in uploaded_files:
                uploaded.seek(0)

            docx_files = read_uploaded_files(uploaded_files)

            if not docx_files:
                st.error("No DOCX files found. Please upload DOCX files or a ZIP containing DOCX files.")
            else:
                for file_name, file_bytes in docx_files:
                    try:
                        ksa_items = extract_ksas_from_docx(file_name, file_bytes)
                        all_ksa_items.extend(ksa_items)
                    except Exception as e:
                        st.warning(f"Failed to extract KSAs from {file_name}: {e}")

                if not all_ksa_items:
                    st.error("No KSA table found in the uploaded documents.")
                else:
                    ksa_df, over_140_df = build_ksa_dataframe(all_ksa_items)

                    st.success(
                        f"Extracted {len(ksa_df)} unique KSA(s). "
                        f"{len(over_140_df)} KSA(s) are over 140 characters."
                    )

                    st.subheader("KSA Upload Table")
                    st.dataframe(ksa_df, use_container_width=True)

                    st.subheader("KSA Over 140 Characters")
                    st.dataframe(over_140_df, use_container_width=True)

                    ksa_excel_output = create_ksa_excel_output(ksa_df, over_140_df)

                    st.download_button(
                        label="Download KSA Excel",
                        data=ksa_excel_output,
                        file_name="KSA_upload_generated.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

    st.markdown("---")
    st.header("Generate Responsibilities Upload Table")

    st.write(
        "This function extracts Responsibilities tables and maps them to KSAs based on the responsibility numbers "
        "listed in the KSA tables. If a responsibility has multiple mapped KSAs, the responsibility information is "
        "only filled once and the additional KSAs are listed in separate rows."
    )

    if st.button("Generate Responsibilities Table"):
        if not uploaded_files:
            st.error("Please upload at least one DOCX file or one ZIP file first.")
        else:
            all_resp_rows = []

            # Important: reset file pointer before reading again
            for uploaded in uploaded_files:
                uploaded.seek(0)

            docx_files = read_uploaded_files(uploaded_files)

            if not docx_files:
                st.error("No DOCX files found. Please upload DOCX files or a ZIP containing DOCX files.")
            else:
                for file_name, file_bytes in docx_files:
                    try:
                        rows = extract_responsibilities_with_ksa_mapping_from_docx(
                            file_name=file_name,
                            file_bytes=file_bytes,
                            folder_name=folder_name
                        )
                        all_resp_rows.extend(rows)
                    except Exception as e:
                        st.warning(f"Failed to extract responsibilities from {file_name}: {e}")

                if not all_resp_rows:
                    st.error("No responsibilities and KSA mapping tables found.")
                else:
                    resp_df, resp_over_140_df = build_responsibilities_dataframe(all_resp_rows)

                    st.success(
                        f"Generated {len(resp_df)} responsibility/KSA mapping row(s). "
                        f"{len(resp_over_140_df)} row(s) are related to responsibilities with KSA over 140 characters."
                    )

                    st.subheader("Responsibilities Upload Table")
                    st.dataframe(resp_df, use_container_width=True)

                    st.subheader("Responsibilities with KSA Over 140 Characters")
                    st.dataframe(resp_over_140_df, use_container_width=True)

                    resp_excel_output = create_responsibilities_excel_output(
                        resp_df,
                        resp_over_140_df
                    )

                    st.download_button(
                        label="Download Responsibilities Excel",
                        data=resp_excel_output,
                        file_name="Responsibilities_upload_generated.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

    st.markdown("---")
    st.header("Generate Specialties Upload Table")

    st.write(
        "This function extracts specialty names, introduction descriptions, mapped responsibilities, "
        "and whether each responsibility is mandatory based on bold formatting in the Word document."
    )

    if st.button("Generate Specialties Table"):
        if not uploaded_files:
            st.error("Please upload at least one DOCX file or one ZIP file first.")
        else:
            all_specialty_rows = []

            # Important: reset file pointer before reading again
            for uploaded in uploaded_files:
                uploaded.seek(0)

            docx_files = read_uploaded_files(uploaded_files)

            if not docx_files:
                st.error("No DOCX files found. Please upload DOCX files or a ZIP containing DOCX files.")
            else:
                for file_name, file_bytes in docx_files:
                    try:
                        rows = extract_specialties_upload_rows_from_docx(
                            file_name=file_name,
                            file_bytes=file_bytes,
                            folder_name=folder_name
                        )
                        all_specialty_rows.extend(rows)
                    except Exception as e:
                        st.warning(f"Failed to extract specialties from {file_name}: {e}")

                if not all_specialty_rows:
                    st.error("No specialties or responsibility tables found.")
                else:
                    specialties_df = build_specialties_dataframe(all_specialty_rows)

                    st.success(f"Generated {len(specialties_df)} specialty/responsibility mapping row(s).")

                    st.subheader("Specialties Upload Table")
                    st.dataframe(specialties_df, use_container_width=True)

                    specialties_excel_output = create_specialties_excel_output(specialties_df)

                    st.download_button(
                        label="Download Specialties Excel",
                        data=specialties_excel_output,
                        file_name="Specialties_upload_generated.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )



def main():
    st.set_page_config(
        page_title="GJP Automation Tool",
        page_icon="📄",
        layout="wide"
    )

    st.sidebar.title("GJP Automation Tool")

    page = st.sidebar.radio(
        "Choose a function",
        [
            "Word to Excel",
            "Word to AoW/KSA Matrix",
            "Excel + Word to New GJP",
            "ksa_conversion_tables",
        ],
        index=2
    )

    if page == "Word to Excel":
        render_word_to_excel()

    elif page == "Word to AoW/KSA Matrix":
        render_aow_ksa_mapping_matrix()

    elif page == "Excel + Word to New GJP":
        render_gjp_document_generator()

    elif page == "ksa_conversion_tables":
        render_ksa_conversion_tables()


if __name__ == "__main__":
    main()
